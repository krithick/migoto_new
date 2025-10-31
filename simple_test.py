#!/usr/bin/env python3
"""
Simple test to check DOCX extraction without dependencies
"""

import asyncio
import json
import io
import docx
from enhanced_scenario_generator import FlexibleScenarioGenerator

async def extract_text_from_docx(file_content):
    """Extract text from .docx file content"""
    try:
        with io.BytesIO(file_content) as f:
            doc = docx.Document(f)
            text_parts = []
            
            print(f"DOCX Document has {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables")
            
            # Extract paragraphs with section markers
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    text_parts.append(f"PARAGRAPH_{i}: {para.text.strip()}")
            
            # Extract tables with better structure preservation
            for table_idx, table in enumerate(doc.tables):
                text_parts.append(f"\nTABLE_{table_idx}_START:")
                for row_idx, row in enumerate(table.rows):
                    row_text = []
                    for cell_idx, cell in enumerate(row.cells):
                        if cell.text.strip():
                            row_text.append(f"CELL_{cell_idx}: {cell.text.strip()}")
                    if row_text:
                        text_parts.append(f"ROW_{row_idx}: {' | '.join(row_text)}")
                text_parts.append(f"TABLE_{table_idx}_END\n")
            
            full_text = "\n".join(text_parts)
            print(f"Extracted {len(full_text)} characters from DOCX")
            
            return full_text if full_text.strip() else None
            
    except Exception as e:
        print(f"DOCX extraction error: {str(e)}")
        return None

async def test_docx_extraction():
    """Test DOCX extraction with the Leadership document"""
    
    print("TESTING DOCX EXTRACTION")
    print("=" * 30)
    
    # Read DOCX file
    with open("Leadership Fundamentals and Styles.docx", "rb") as f:
        docx_content = f.read()
    
    # Extract text
    document_content = await extract_text_from_docx(docx_content)
    
    if not document_content:
        print("Failed to extract content from DOCX")
        return
    
    print(f"\nFirst 1000 characters of extracted content:")
    print("-" * 50)
    print(document_content[:1000])
    print("-" * 50)
    
    # Test with enhanced generator
    generator = FlexibleScenarioGenerator(client=None)
    
    print("\nTesting raw data extraction...")
    raw_data = generator._extract_raw_data(document_content)
    
    print("\nTesting template building...")
    template = generator._build_template_from_raw_data(raw_data)
    
    # Save results
    with open("docx_test_results.json", "w", encoding="utf-8") as f:
        json.dump({
            "extracted_text_sample": document_content[:2000],
            "raw_data": raw_data,
            "template": template
        }, f, indent=2, ensure_ascii=False)
    
    print(f"\nResults saved to: docx_test_results.json")
    print("Test completed!")

if __name__ == "__main__":
    asyncio.run(test_docx_extraction())