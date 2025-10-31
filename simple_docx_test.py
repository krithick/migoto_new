import io
import docx
import os

def test_docx_extraction():
    try:
        # Read the DOCX file
        docx_path = "Leadership Fundamentals and Styles.docx"
        if not os.path.exists(docx_path):
            print(f"File not found: {docx_path}")
            return
            
        with open(docx_path, 'rb') as f:
            content = f.read()
        
        print(f"Testing DOCX extraction on {docx_path}")
        print(f"File size: {len(content)} bytes")
        
        # Extract using docx library
        with io.BytesIO(content) as f:
            doc = docx.Document(f)
            text_parts = []
            
            print(f"Document has {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables")
            
            # Extract paragraphs with section markers
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    text_parts.append(f"PARAGRAPH_{i}: {para.text.strip()}")
            
            # Extract tables with better structure preservation
            for table_idx, table in enumerate(doc.tables):
                text_parts.append(f"\\nTABLE_{table_idx}_START:")
                for row_idx, row in enumerate(table.rows):
                    row_text = []
                    for cell_idx, cell in enumerate(row.cells):
                        if cell.text.strip():
                            row_text.append(f"CELL_{cell_idx}: {cell.text.strip()}")
                    if row_text:
                        text_parts.append(f"ROW_{row_idx}: {' | '.join(row_text)}")
                text_parts.append(f"TABLE_{table_idx}_END\\n")
            
            full_text = "\\n".join(text_parts)
            
        print(f"Extracted {len(full_text)} characters")
        
        # Analyze structure
        table_count = full_text.count('TABLE_')
        row_count = full_text.count('ROW_')
        cell_count = full_text.count('CELL_')
        
        print(f"Structure Analysis:")
        print(f"  Tables found: {table_count}")
        print(f"  Rows found: {row_count}")
        print(f"  Cells found: {cell_count}")
        
        # Check for key sections
        sections = []
        for i in range(10):
            if f'TABLE_{i}_START:' in full_text:
                sections.append(f'TABLE_{i}')
        
        print(f"Sections found: {sections}")
        
        # Check for key content
        key_checks = {
            "Course info": "Course" in full_text,
            "Leadership": "leadership" in full_text.lower(),
            "AI Trainer": "AI Trainer" in full_text,
            "Success Metrics": "Success Metrics" in full_text,
            "Common Mistakes": "Common Mistake" in full_text
        }
        
        print("Key content check:")
        for content_type, found in key_checks.items():
            status = "YES" if found else "NO"
            print(f"  {content_type}: {status}")
        
        # Show sample
        print("\\nSample content (first 1000 chars):")
        print("-" * 50)
        print(full_text[:1000])
        print("-" * 50)
        
        # Quality score
        quality_score = len(sections) * 10 + sum(key_checks.values()) * 5
        print(f"\\nExtraction Quality Score: {quality_score}/125")
        
        if quality_score >= 100:
            print("EXCELLENT QUALITY - Ready for production!")
        elif quality_score >= 75:
            print("GOOD QUALITY - Minor improvements needed")
        elif quality_score >= 50:
            print("MODERATE QUALITY - Some issues to fix")
        else:
            print("POOR QUALITY - Major fixes needed")
            
    except Exception as e:
        print(f"ERROR: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    test_docx_extraction()