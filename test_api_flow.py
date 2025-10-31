#!/usr/bin/env python3
"""
Test the full API flow: Document → Extract → Review → Generate
"""

import asyncio
import json
from enhanced_scenario_generator import FlexibleScenarioGenerator

async def test_full_flow():
    """Test the complete extraction → validation → prompt generation flow"""
    
    print("TESTING FULL API FLOW")
    print("=" * 30)
    
    # Step 1: Extract from DOCX (simulating API call)
    print("\n1. Document Analysis...")
    
    # Read DOCX and extract
    import io
    import docx
    
    async def extract_text_from_docx(file_content):
        with io.BytesIO(file_content) as f:
            doc = docx.Document(f)
            text_parts = []
            
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    text_parts.append(f"PARAGRAPH_{i}: {para.text.strip()}")
            
            for table_idx, table in enumerate(doc.tables):
                text_parts.append(f"\nTABLE_{table_idx}_START:")
                for row_idx, row in enumerate(table.rows):
                    row_text = []
                    for cell_idx, cell in enumerate(row.cells):
                        if cell.text.strip():
                            row_text.append(f"CELL_{cell_idx}: {cell.text.strip()}")
                    if row_text:
                        text_parts.append(f"ROW_{row_idx}: {' | '.join(row_text)}")
                text_parts.append(f"TABLE_{table_idx}_END\n")
            
            return "\n".join(text_parts)
    
    with open("Leadership Fundamentals and Styles.docx", "rb") as f:
        docx_content = f.read()
    
    document_text = await extract_text_from_docx(docx_content)
    
    # Extract using enhanced system
    generator = FlexibleScenarioGenerator(client=None)  # No LLM for testing
    extracted_data = await generator.flexible_extract_from_document(document_text)
    
    print(f"Extracted data with {len(extracted_data)} sections")
    print(f"   Expert Role: {extracted_data.get('participant_roles', {}).get('expert_role', 'Not found')[:50]}...")
    print(f"   Practice Role: {extracted_data.get('participant_roles', {}).get('practice_role', 'Not found')[:50]}...")
    
    # Step 2: Validate and enhance (simulating user review)
    print("\n2. Template Validation...")
    
    validation_result = await generator.validate_and_enhance_template(extracted_data)
    
    print(f"Validation completed")
    print(f"   Completeness Score: {validation_result['validation_notes']['completeness_score']}%")
    print(f"   Missing Elements: {len(validation_result['validation_notes']['missing_elements'])}")
    print(f"   Suggestions: {len(validation_result['validation_notes']['suggestions'])}")
    
    # Step 3: Generate prompts (simulating final generation)
    print("\n3. Prompt Generation...")
    
    scenario_data = validation_result['validated_extraction']
    
    # Generate learn mode prompt
    expert_role = scenario_data.get('participant_roles', {}).get('expert_role', 'Expert Trainer')
    topics = scenario_data.get('content_specifics', {}).get('key_knowledge', ['General topics'])
    learn_prompt = f"You are a {expert_role}. Your role is to teach about: {', '.join(topics[:3])}. Maintain a supportive and educational approach."
    
    # Generate assess mode prompt  
    practice_role = scenario_data.get('participant_roles', {}).get('practice_role', 'Practice Partner')
    concerns = scenario_data.get('conversation_dynamics', {}).get('typical_interactions', ['General questions'])
    assess_prompt = f"You are playing the role of: {practice_role}. Your typical concerns include: {', '.join(concerns[:2])}. Engage naturally with the learner."
    
    print(f"Prompts generated")
    print(f"   Learn Mode: {learn_prompt[:100]}...")
    print(f"   Assess Mode: {assess_prompt[:100]}...")
    
    # Step 4: Save results
    results = {
        "extraction_results": {
            "completeness_score": validation_result['validation_notes']['completeness_score'],
            "missing_elements": validation_result['validation_notes']['missing_elements'],
            "key_data": {
                "main_topic": scenario_data.get('scenario_understanding', {}).get('main_topic'),
                "expert_role": scenario_data.get('participant_roles', {}).get('expert_role'),
                "practice_role": scenario_data.get('participant_roles', {}).get('practice_role'),
                "target_skills_count": len(scenario_data.get('scenario_understanding', {}).get('target_skills', [])),
                "knowledge_dos_count": len(scenario_data.get('knowledge_base', {}).get('dos', [])),
                "knowledge_donts_count": len(scenario_data.get('knowledge_base', {}).get('donts', []))
            }
        },
        "generated_prompts": {
            "learn_mode": learn_prompt,
            "assess_mode": assess_prompt,
            "try_mode": assess_prompt  # Same as assess for now
        }
    }
    
    with open("api_flow_test_results.json", "w", encoding="utf-8") as f:
        json.dump(results, f, indent=2, ensure_ascii=False)
    
    print(f"\nFULL FLOW TEST COMPLETED")
    print(f"   Results saved to: api_flow_test_results.json")
    print(f"   Ready for API integration!")

if __name__ == "__main__":
    asyncio.run(test_full_flow())