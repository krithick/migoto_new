from fastapi import APIRouter, Depends, HTTPException, status, Body, Path, Query , UploadFile, Form , File
from typing import List, Optional, Dict, Any
from uuid import UUID
from datetime import datetime
from typing import List, Optional, Union, Dict, Any, Set
from pydantic import BaseModel, Field, EmailStr, validator, root_validator,model_validator
from uuid import UUID, uuid4
from datetime import datetime
from enum import Enum
# from main import azure_openai_client
from dotenv import load_dotenv
from openai import AzureOpenAI ,AsyncAzureOpenAI
import os
from utils import convert_template_to_markdown
from models.user_models import UserRole
load_dotenv(".env")
from core.simple_token_logger import log_token_usage
from core.archetype_classifier import ArchetypeClassifier
from database import db,get_db
from core.scenario_extractor_v2 import ScenarioExtractorV2
router = APIRouter(prefix="/scenario", tags=["Scenario Generation"])
api_key = os.getenv("api_key")
endpoint = os.getenv("endpoint")
api_version =  os.getenv("api_version")
        
azure_openai_client = AsyncAzureOpenAI(
            api_key=api_key,
            api_version=api_version,
            azure_endpoint=endpoint
        )

import io
import docx
import pypdf
from pydantic import BaseModel
from typing import Optional, Dict, Any, List
import json
import re
import uuid
from datetime import datetime
from fastapi import APIRouter, HTTPException, File, UploadFile, Body
from database import db

class PersonaDetails(BaseModel):
    name: str
    description: str
    persona_type: str
    gender: str
    age: int
    character_goal: str
    location: str
    persona_details: str
    situation: str
    context_type: str
    background_story: str
    full_persona: Optional[Dict[str, Any]] = None

class TemplateAnalysisResponse(BaseModel):
    general_info: Dict[str, Any]
    context_overview: Dict[str, Any]
    persona_definitions: Dict[str, Any]
    dialogue_flow: Dict[str, Any]
    knowledge_base: Dict[str, Any]
    variations_challenges: Dict[str, Any]
    success_metrics: Dict[str, Any]
    feedback_mechanism: Dict[str, Any]

class ScenarioResponse(BaseModel):
    learn_mode: str
    assess_mode: str
    try_mode: str
    scenario_title: str
    extracted_info: Optional[Dict[str, Any]] = None
    generated_persona: Optional[Dict[str, Any]] = None
    knowledge_base_id: Optional[str] = None  # 🔑 ADD THIS

class ScenarioListResponse(BaseModel):
    scenarios: List[Dict[str, Any]]
    total_count: int
async def extract_text_from_docx(file_content):
    """Enhanced DOCX extraction with better error handling and table support"""
    try:
        with io.BytesIO(file_content) as f:
            doc = docx.Document(f)
            text_parts = []
            
            print(f"DOCX Document has {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables")
            
            # Extract paragraphs with section markers
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    text_parts.append(f"PARAGRAPH_{i}: {para.text.strip()}")
            
            # Extract tables with better structure preservation
            for table_idx, table in enumerate(doc.tables):
                text_parts.append(f"\nTABLE_{table_idx}_START:")
                for row_idx, row in enumerate(table.rows):
                    row_text = []
                    for cell_idx, cell in enumerate(row.cells):
                        if cell.text.strip():
                            row_text.append(f"CELL_{cell_idx}: {cell.text.strip()}")
                    if row_text:
                        text_parts.append(f"ROW_{row_idx}: {' | '.join(row_text)}")
                text_parts.append(f"TABLE_{table_idx}_END\n")
            
            full_text = "\n".join(text_parts)
            print(f"Extracted {len(full_text)} characters from DOCX")
            print(f"First 500 characters: {full_text[:500]}...")
            
            # Check for key template sections
            key_sections = ["SECTION 1", "SECTION 2", "SECTION 3", "AI Trainer Role", "Success Metrics"]
            found_sections = [section for section in key_sections if section.lower() in full_text.lower()]
            print(f"Found template sections: {found_sections}")
            
            return full_text if full_text.strip() else None
            
    except Exception as e:
        print(f"DOCX extraction error: {str(e)}")
        print(f"Error type: {type(e).__name__}")
        import traceback
        print(f"Full traceback: {traceback.format_exc()}")
        return None


async def extract_text_from_pdf(file_content):
    """Extract text from PDF file content using pypdf"""
    try:
        with io.BytesIO(file_content) as f:
            reader = pypdf.PdfReader(f)
            text = ""
            for page_num in range(len(reader.pages)):
                text += reader.pages[page_num].extract_text()
        return text.strip()
    except Exception as e:
        print(f"Error extracting text from PDF: {str(e)}")
        return None

# API Endpoints

@router.post("/test-archetype-classification")
async def test_archetype_classification(
    scenario_document: str = Body(..., embed=True),
    db: Any = Depends(get_db)
):
    """
    Test endpoint: Classify scenario archetype without full template generation
    """
    try:
        extractor = ScenarioExtractorV2(azure_openai_client, model="gpt-4o")
        
        # Quick extraction for classification
        template_data = await extractor.extract_scenario_info(scenario_document)
        
        archetype_info = template_data.get("archetype_classification", {})
        
        return {
            "scenario_preview": scenario_document[:200] + "...",
            "archetype_classification": archetype_info,
            "scenario_title": template_data.get("context_overview", {}).get("scenario_title", "Unknown"),
            "domain": template_data.get("general_info", {}).get("domain", "Unknown"),
            "message": "Archetype classification completed"
        }
    except Exception as e:
        print(f"Error in test_archetype_classification: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/analyze-scenario", response_model=TemplateAnalysisResponse)
async def analyze_scenario_endpoint(scenario_document: str = Body(..., embed=True)):
    """
    Step 1: Analyze raw scenario description and create comprehensive template structure.
    """
    try:
        # Initialize V2 extractor
        extractor = ScenarioExtractorV2(azure_openai_client, model="gpt-4o")
        template_data = await extractor.extract_scenario_info(scenario_document)
        evaluation_metrics = await extractor.extract_evaluation_metrics_from_template(scenario_document, template_data)
        template_data["evaluation_metrics"] = evaluation_metrics        
        return TemplateAnalysisResponse(
            general_info=template_data.get("general_info", {}),
            context_overview=template_data.get("context_overview", {}),
            persona_definitions=template_data.get("persona_definitions", {}),
            dialogue_flow=template_data.get("dialogue_flow", {}),
            knowledge_base=template_data.get("knowledge_base", {}),
            variations_challenges=template_data.get("variations_challenges", {}),
            success_metrics=template_data.get("success_metrics", {}),
            feedback_mechanism=template_data.get("feedback_mechanism", {})
        )
    except Exception as e:
        print(f"Error in analyze_scenario_endpoint: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/edit-template-section")
async def edit_template_section(
    section_name: str = Body(...),
    section_data: Dict[str, Any] = Body(...),
    current_template: Dict[str, Any] = Body(...)
):
    """
    Step 2: Edit specific sections of the template.
    """
    try:
        # Update the specific section in the template
        if section_name in current_template:
            current_template[section_name].update(section_data)
        else:
            current_template[section_name] = section_data
        
        return {
            "message": f"Section '{section_name}' updated successfully",
            "updated_template": current_template,
            "section_preview": current_template.get(section_name, {})
        }
    except Exception as e:
        print(f"Error in edit_template_section: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/add-remove-points")
async def add_remove_points(
    section: str = Body(...),
    subsection: str = Body(...),
    action: str = Body(...),
    point_data: Dict[str, Any] = Body(...),
    current_template: Dict[str, Any] = Body(...)
):
    """
    Step 2.1: Add, remove, or edit specific points in knowledge base or dialogue flow.
    """
    try:
        if section not in current_template:
            current_template[section] = {}
        
        if subsection not in current_template[section]:
            current_template[section][subsection] = []
        
        points_list = current_template[section][subsection]
        
        if action == "add":
            new_point = point_data.get("content", "")
            if new_point:
                points_list.append(new_point)
        
        elif action == "remove":
            index = point_data.get("index", -1)
            if 0 <= index < len(points_list):
                points_list.pop(index)
            else:
                raise HTTPException(status_code=400, detail="Invalid index")
        
        elif action == "edit":
            index = point_data.get("index", -1)
            new_content = point_data.get("content", "")
            if 0 <= index < len(points_list) and new_content:
                points_list[index] = new_content
            else:
                raise HTTPException(status_code=400, detail="Invalid index or content")
        
        current_template[section][subsection] = points_list
        
        return {
            "message": f"Successfully {action}ed point in {section}.{subsection}",
            "updated_points": points_list,
            "updated_template": current_template
        }
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error in add_remove_points: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/generate-final-prompts", response_model=ScenarioResponse)
async def generate_final_prompts(template_data: Dict[str, Any] = Body(...)):
    """
    DEPRECATED: Use /generate-prompts-for-persona instead for better control.
    
    This generates ONE set of prompts with ONE default persona.
    For multiple personas, use the new workflow.
    """
    try:
        from core.system_prompt_generator import SystemPromptGenerator
        from core.learn_mode_prompt_generator import LearnModePromptGenerator
        
        generator = EnhancedScenarioGenerator(azure_openai_client)
        
        # Generate ONE default persona
        personas = await generator.generate_personas_from_template(template_data)
        
        # Initialize new prompt generators
        system_prompt_gen = SystemPromptGenerator(azure_openai_client, model="gpt-4o")
        learn_prompt_gen = LearnModePromptGenerator(azure_openai_client, model="gpt-4o")
        
        # Generate learn mode prompt (trainer, no persona needed)
        learn_mode_prompt = await learn_prompt_gen.generate_trainer_prompt(template_data)
        
        # Generate assess/try mode prompts (SAME prompt for both)
        assess_persona = personas.get("assess_mode_character", {})
        
        assess_try_prompt = await system_prompt_gen.generate_system_prompt(
            template_data=template_data,
            persona_data=assess_persona,
            mode="assess_mode"  # Same for both assess and try
        )
        
        # Get knowledge base ID
        template_id = template_data.get("template_id")
        knowledge_base_id = None
        if template_id:
            template_record = await db.templates.find_one({"id": template_id})
            if template_record:
                knowledge_base_id = template_record.get("knowledge_base_id")
        
        return ScenarioResponse(
            learn_mode=learn_mode_prompt,
            assess_mode=assess_try_prompt,  # Same prompt
            try_mode=assess_try_prompt,     # Same prompt
            scenario_title=template_data.get("context_overview", {}).get("scenario_title", "Training Scenario"),
            extracted_info=template_data,
            generated_persona=personas,
            knowledge_base_id=knowledge_base_id
        )
    except Exception as e:
        print(f"Error in generate_final_prompts: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/quick-generate-prompts", response_model=ScenarioResponse)
async def quick_generate_prompts(scenario_document: str = Body(..., embed=True)):
    """
    Quick Generation: All steps in one API call.
    """
    try:
        generator = EnhancedScenarioGenerator(azure_openai_client)
        
        # Step 1: Analyze scenario to template
        template_data = await generator.extract_scenario_info(scenario_document)
        
        # Step 3: Generate final prompts
        personas = await generator.generate_personas_from_template(template_data)
        
        learn_mode_prompt = await generator.generate_learn_mode_from_template(template_data)
        assess_mode_prompt = await generator.generate_assess_mode_from_template(template_data)
        try_mode_prompt = await generator.generate_try_mode_from_template(template_data)
        
        # Apply personas and language instructions
        learn_mode_prompt = generator.insert_persona(learn_mode_prompt, personas.get("learn_mode_expert", {}))
        assess_mode_prompt = generator.insert_persona(assess_mode_prompt, personas.get("assess_mode_character", {}))
        try_mode_prompt = generator.insert_persona(try_mode_prompt, personas.get("assess_mode_character", {}))
        
        language_data = template_data.get("general_info", {})
        learn_mode_prompt = generator.insert_language_instructions(learn_mode_prompt, language_data)
        assess_mode_prompt = generator.insert_language_instructions(assess_mode_prompt, language_data)
        try_mode_prompt = generator.insert_language_instructions(try_mode_prompt, language_data)
        
        return ScenarioResponse(
            learn_mode=learn_mode_prompt,
            assess_mode=assess_mode_prompt,
            try_mode=try_mode_prompt,
            scenario_title=template_data.get("context_overview", {}).get("scenario_title", "Training Scenario"),
            extracted_info=template_data,
            generated_persona=personas
        )
    except Exception as e:
        print(f"Error in quick_generate_prompts: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/regenerate-prompts", response_model=ScenarioResponse)
async def regenerate_prompts(template_data: Dict[str, Any] = Body(...)):
    """
    Regenerate prompts from existing template data after user edits.
    """
    try:
        generator = EnhancedScenarioGenerator(azure_openai_client)
        
        # Use existing personas or generate new ones if needed
        personas = template_data.get("generated_personas")
        if not personas:
            personas = await generator.generate_personas_from_template(template_data)
        
        # Generate all three prompts using updated template data
        learn_mode_prompt = await generator.generate_learn_mode_from_template(template_data)
        assess_mode_prompt = await generator.generate_assess_mode_from_template(template_data)
        try_mode_prompt = await generator.generate_try_mode_from_template(template_data)
        
        # Apply personas and language instructions
        learn_mode_prompt = generator.insert_persona(learn_mode_prompt, personas.get("learn_mode_expert", {}))
        assess_mode_prompt = generator.insert_persona(assess_mode_prompt, personas.get("assess_mode_character", {}))
        try_mode_prompt = generator.insert_persona(try_mode_prompt, personas.get("assess_mode_character", {}))
        
        language_data = template_data.get("general_info", {})
        learn_mode_prompt = generator.insert_language_instructions(learn_mode_prompt, language_data)
        assess_mode_prompt = generator.insert_language_instructions(assess_mode_prompt, language_data)
        try_mode_prompt = generator.insert_language_instructions(try_mode_prompt, language_data)
        
        return ScenarioResponse(
            learn_mode=learn_mode_prompt,
            assess_mode=assess_mode_prompt,
            try_mode=try_mode_prompt,
            scenario_title=template_data.get("context_overview", {}).get("scenario_title", "Training Scenario"),
            extracted_info=template_data,
            generated_persona=personas
        )
    except Exception as e:
        print(f"Error in regenerate_prompts: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/save-template-to-db")
async def save_template_to_db(
    template_data: Dict[str, Any] = Body(...),
    template_name: str = Body(...),
    template_id: Optional[str] = Body(default=None)
):
    """Save template data to database"""
    try:
        template_record = {
            "id": template_id or str(uuid.uuid4()),
            "name": template_name,
            "template_data": template_data,
            "created_at": datetime.now().isoformat(),
            "updated_at": datetime.now().isoformat()
        }
        
        # Save to your database
        db.templates.insert_one(template_record)  # or update if template_id exists
        
        return {
            "message": "Template saved to database successfully",
            "template_id": template_record["id"]
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
def convert_objectid(document):
    """Convert ObjectId to str in a MongoDB document"""
    if document and "_id" in document:
        document["_id"] = str(document["_id"])
    return document
@router.get("/load-template-from-db/{template_id}")
async def load_template_from_db(template_id: str):
    """Load template data from database"""
    try:
        template = await db.templates.find_one({"id": template_id})
        print(template,"template")
        if not template:
            raise HTTPException(status_code=404, detail="Template not found")
        
        return convert_objectid(template)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/list-templates-from-db")
async def list_templates_from_db():
    try:
        cursor = db.templates.find({}, {"_id": 0})
        templates = await cursor.to_list(length=100)
        return {"templates": templates}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
@router.put("/update-template-in-db/{template_id}")
async def update_template_in_db(
    template_id: str,
    template_data: Dict[str, Any] = Body(...),
    template_name: str = Body(...)
):
    """Update existing template in database"""
    try:
        # Check if template exists
        existing_template = db.templates.find_one({"id": template_id})
        if not existing_template:
            raise HTTPException(status_code=404, detail="Template not found")
        
        # Update template
        updated_template = {
            "name": template_name,
            "template_data": template_data,
            "updated_at": datetime.now().isoformat()
        }
        
        db.templates.update_one(
            {"id": template_id}, 
            {"$set": updated_template}
        )
        
        return {
            "message": "Template updated successfully",
            "template_id": template_id
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.delete("/delete-template-from-db/{template_id}")
async def delete_template_from_db(template_id: str):
    """Delete template from database"""
    try:
        # Check if template exists
        existing_template = db.templates.find_one({"id": template_id})
        if not existing_template:
            raise HTTPException(status_code=404, detail="Template not found")
        
        # Delete template
        result = db.templates.delete_one({"id": template_id})
        
        if result.deleted_count == 0:
            raise HTTPException(status_code=404, detail="Template not found")
        
        return {
            "message": "Template deleted successfully",
            "deleted_template": {
                "id": template_id,
                "name": existing_template.get("name")
            }
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
# Additional utility endpoints

@router.post("/generate-personas-from-template")
async def generate_personas_from_template_endpoint(template_data: Dict[str, Any] = Body(...)):
    """
    Generate detailed personas based on template persona definitions.
    """
    try:
        generator = EnhancedScenarioGenerator(azure_openai_client)
        personas = await generator.generate_personas_from_template(template_data)
        
        return {
            "learn_mode_expert": personas.get("learn_mode_expert", {}),
            "assess_mode_character": personas.get("assess_mode_character", {}),
            "scenario_info": {
                "title": template_data.get("context_overview", {}).get("scenario_title", "Training Scenario"),
                "domain": template_data.get("general_info", {}).get("domain", "general"),
                "purpose": template_data.get("context_overview", {}).get("purpose_of_scenario", "learning scenario")
            }
        }
    except Exception as e:
        print(f"Error in generate_personas_from_template_endpoint: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/file-to-template")
async def file_to_template(file: UploadFile = File(...)):
    try:
        print(f"Processing file: {file.filename}")
        print(f"Content type: {file.content_type}")
        
        content = await file.read()
        print(f"File size: {len(content)} bytes")
        
        # Process based on file type
        scenario_text = ""
        extraction_method = "unknown"
        
        if file.filename.lower().endswith('.txt'):
            scenario_text = content.decode('utf-8')
            extraction_method = "text"
            
        elif file.filename.lower().endswith(('.doc', '.docx')):
            scenario_text = await extract_text_from_docx(content)
            extraction_method = "docx"
            
            if scenario_text is None:
                # Try alternative extraction
                print("Primary DOCX extraction failed, trying alternative...")
                try:
                    scenario_text = content.decode('utf-8', errors='ignore')
                    extraction_method = "fallback_text"
                except:
                    raise HTTPException(400, "Could not extract text from DOCX file")
        
        # Validate extraction
        if not scenario_text or len(scenario_text.strip()) < 50:
            raise HTTPException(400, f"Insufficient content extracted. Got {len(scenario_text)} characters")
        
        print(f"Successfully extracted {len(scenario_text)} characters using {extraction_method}")
        print(f"Content preview: {scenario_text[:300]}...")
        
        # Process with LLM
        generator = EnhancedScenarioGenerator(azure_openai_client)
        
        # Check if client is properly initialized
        if generator.client is None:
            print("WARNING: LLM client not initialized, using enhanced mock data")
            template_data = generator._get_enhanced_mock_template_data(scenario_text)
        else:
            template_data = await generator.extract_scenario_info(scenario_text)
        
        return template_data
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"File processing error: {str(e)}")
        print(f"Error type: {type(e).__name__}")
        raise HTTPException(500, f"File processing failed: {str(e)}")
@router.get("/template-schema")
async def get_template_schema():
    """
    Get the complete template schema for UI builders.
    """
    try:
        return {
            "general_info": {
                "domain": "string - Healthcare/Banking/Sales/Education/etc.",
                "purpose_of_ai_bot": "string - Trainer/Customer/Student/etc.", 
                "target_audience": "string - Employees/General Public/etc.",
                "preferred_language": "string - English/Spanish/etc."
            },
            "context_overview": {
                "scenario_title": "string - Descriptive title",
                "learn_mode_description": "string - What happens in learn mode",
                "assess_mode_description": "string - What happens in assessment mode",
                "try_mode_description": "string - What happens in try mode", 
                "purpose_of_scenario": "string - Learning objectives"
            },
            "knowledge_base": {
                "dos": "array - List of best practices",
                "donts": "array - List of things to avoid",
                "key_facts": "array - Important facts about the domain",
                "conversation_topics": "array - Topics to explore in conversation"
            },
            "feedback_mechanism": {
                "positive_closing": "string - Positive ending message",
                "negative_closing": "string - Negative ending message",
                "neutral_closing": "string - Neutral ending message"
            }
        }
    except Exception as e:
        print(f"Error in get_template_schema: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))
# ADD these imports at the top of your existing file
from core.document_processor import DocumentProcessor
from core.azure_search_manager import AzureVectorSearchManager
from core.user import get_current_user
from models.user_models import UserDB
from database import get_db

# newwww
@router.post("/analyze-template-with-docs")
async def analyze_template_with_docs(
    template_file: UploadFile = File(...),
    supporting_docs: List[UploadFile] = File(...),
    template_name: str = Form(...),
    current_user: UserDB = Depends(get_current_user),
    db: Any = Depends(get_db)
):
    """Step 1: Analyze template + process docs, return editable template"""
    try:
        # 1. Process template file
        content = await template_file.read()
        
        if template_file.filename.lower().endswith(('.doc', '.docx')):
            scenario_text = await extract_text_from_docx(content)
            
            # Fallback extraction if primary method fails
            if not scenario_text or len(scenario_text.strip()) < 100:
                print("Primary DOCX extraction failed, trying alternative methods...")
                try:
                    # Try simple text extraction
                    scenario_text = content.decode('utf-8', errors='ignore')
                    print(f"Fallback text extraction got {len(scenario_text)} characters")
                except Exception as e2:
                    print(f"Fallback extraction also failed: {e2}")
                    raise HTTPException(400, "Could not extract text from Word document")
                    
        elif template_file.filename.lower().endswith('.pdf'):
            scenario_text = await extract_text_from_pdf(content)
        else:
            scenario_text = content.decode('utf-8')
            
        # Validate extraction before processing
        if not scenario_text or len(scenario_text.strip()) < 50:
            raise HTTPException(400, f"Insufficient content extracted from document. Got {len(scenario_text) if scenario_text else 0} characters")
            
        print(f"Successfully extracted {len(scenario_text)} characters for processing")
        print(f"Content preview: {scenario_text[:300]}...")
            
        # 2. Analyze template to get editable structure
        generator = EnhancedScenarioGenerator(azure_openai_client)
        template_data = await generator.extract_scenario_info(scenario_text)
        evaluation_metrics = await generator.extract_evaluation_metrics_from_template(scenario_text, template_data)
        template_data["evaluation_metrics"] = evaluation_metrics
        # 3. Process supporting documents → Create knowledge base
        knowledge_base_id = f"kb_{str(uuid4())}"
        supporting_docs_metadata = []
        
        if supporting_docs and len(supporting_docs) > 0:
            supporting_docs_metadata = await process_and_index_documents(
                supporting_docs, knowledge_base_id, db
            )
        
        # 4. Store template for editing (NO prompts yet)
        template_record = {
            "id": str(uuid4()),
            "name": template_name,
            "template_data": template_data,  # EDITABLE
            # "evaluation_metrics": evaluation_metrics,
            "knowledge_base_id": knowledge_base_id if supporting_docs_metadata else None,
            "supporting_documents": len(supporting_docs_metadata),
            "status": "ready_for_editing",
            "created_at": datetime.now().isoformat(),
            "updated_at": datetime.now().isoformat()
        }
        
        await db.templates.insert_one(template_record)
        
        # 5. Create knowledge base record
        if supporting_docs_metadata:
            knowledge_base_record = {
                "_id": knowledge_base_id,
                "template_id": template_record["id"],
                "scenario_title": template_data.get("context_overview", {}).get("scenario_title", template_name),
                "supporting_documents": supporting_docs_metadata,
                "total_documents": len(supporting_docs_metadata),
                "total_chunks": sum(doc.get("chunk_count", 0) for doc in supporting_docs_metadata),
                "created_at": datetime.now(),
                "last_updated": datetime.now(),
                "fact_checking_enabled": True
            }
            await db.knowledge_bases.insert_one(knowledge_base_record)
        
        return {
            "template_id": template_record["id"],
            "template_data": template_data,  # For frontend editing
            # "evaluation_metrics": evaluation_metrics,
            "knowledge_base_id": knowledge_base_id if supporting_docs_metadata else None,
            "supporting_documents_count": len(supporting_docs_metadata),
            "message": "Template analyzed and ready for editing",
            "next_step": "Edit template sections, then generate prompts"
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error: {str(e)}")

# newww
# ADD this helper function
async def process_and_index_documents(supporting_docs: List[UploadFile], knowledge_base_id: str, db: Any) -> List[dict]:
    """Process documents and index in Azure Search"""
    vector_search = AzureVectorSearchManager()
    
    # Create index if it doesn't exist
    try:
        from core.azure_search_setup import AzureSearchIndexManager
        index_manager = AzureSearchIndexManager()
        
        # Check if index exists, create if not
        if not await index_manager.index_exists():
            print("Creating Azure Search index...")
            await index_manager.create_knowledge_base_index()
            print("✅ Index created successfully")
    except Exception as e:
        print(f"⚠️  Index creation warning: {e}")
        
    documents_metadata = []
    
    from openai import AsyncAzureOpenAI
    import os
    
    openai_client = AsyncAzureOpenAI(
        api_key=os.getenv("api_key"),
        azure_endpoint=os.getenv("endpoint"),
        api_version=os.getenv("api_version")
    )
    
    document_processor = DocumentProcessor(openai_client, db)
    vector_search = AzureVectorSearchManager()
    
    for doc_file in supporting_docs:
        try:
            content = await doc_file.read()
            doc_id = str(uuid4())
            
            # Process document (extract text, create chunks, embeddings)
            chunks = await document_processor.process_document(doc_id, content, doc_file.filename)
            
            # Set knowledge base for chunks
            for chunk in chunks:
                chunk.knowledge_base_id = knowledge_base_id
            
            # Index in Azure Search
            await vector_search.index_document_chunks(chunks, knowledge_base_id)
            
            # Store metadata
            # doc_metadata = {
            #     "_id": doc_id,
            #     "knowledge_base_id": knowledge_base_id,
            #     "filename": doc_file.filename,
            #     "file_size": len(content),
            #     "chunk_count": len(chunks),
            #     "processed_at": datetime.now().isoformat()
            # }
            doc_metadata = {
                "_id": doc_id,
                "knowledge_base_id": knowledge_base_id,
                "filename": doc_file.filename,
                "original_filename": doc_file.filename,
                "file_size": len(content),
                "content_type": doc_file.content_type,
                "processing_status": "completed",
                "chunk_count": len(chunks),
                "processed_at": datetime.now().isoformat(),
                "indexed_in_vector_search": True
            }            
            await db.supporting_documents.insert_one(doc_metadata)
            documents_metadata.append(doc_metadata)
            
        except Exception as e:
            print(f"Error processing {doc_file.filename}: {e}")
    
    return documents_metadata

@router.put("/template/{template_id}/evaluation-metrics")
async def update_evaluation_metrics(
    template_id: str,
    evaluation_metrics: Dict[str, Any] = Body(...),
    current_user: UserDB = Depends(get_current_user),
    db: Any = Depends(get_db)
):
    """Edit evaluation metrics for a template"""
    try:
        # Check if template exists
        template = await db.templates.find_one({"id": template_id})
        if not template:
            raise HTTPException(status_code=404, detail="Template not found")
        
        # Update evaluation metrics
        await db.templates.update_one(
            {"id": template_id},
            {"$set": {
                "evaluation_metrics": evaluation_metrics,
                "updated_at": datetime.now().isoformat()
            }}
        )
        
        return {
            "message": "Evaluation metrics updated successfully",
            "template_id": template_id,
            "updated_metrics": evaluation_metrics
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/template/{template_id}/evaluation-metrics")
async def get_evaluation_metrics(
    template_id: str,
    current_user: UserDB = Depends(get_current_user),
    db: Any = Depends(get_db)
):
    """Get evaluation metrics for a template"""
    try:
        template = await db.templates.find_one({"id": template_id})
        if not template:
            raise HTTPException(status_code=404, detail="Template not found")
        
        return {
            "template_id": template_id,
            "evaluation_metrics": template.get("evaluation_metrics", {}),
            "template_name": template.get("name", "")
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
@router.post("/generate-personas")
async def generate_personas(
    template_id: str = Body(...),
    persona_type: str = Body(..., description="learn_mode_expert or assess_mode_character"),
    gender: str = Body(default="",description="Specify Gender if needed"),
    prompt: str = Body(default="",description="Specify context if needed"),
    count: int = Body(default=1, description="Number of personas to generate"),
    db: Any = Depends(get_db)
):
    """
    Generate personas from template ID using PersonaGeneratorV2.
    Returns personas that can be used to generate prompts.
    """
    try:
        from core.persona_generator_v2 import PersonaGeneratorV2
        
        # Get template data from database
        template = await db.templates.find_one({"id": template_id})
        if not template:
            raise HTTPException(status_code=404, detail="Template not found")
        
        template_data = template.get("template_data", {})
        
        # Map persona_type to mode
        mode_map = {
            "learn_mode_expert": "learn_mode",
            "assess_mode_character": "assess_mode",
            "try_mode_character": "try_mode"
        }
        mode = mode_map.get(persona_type, "assess_mode")
        
        # Generate V2 persona
        generator_v2 = PersonaGeneratorV2(azure_openai_client)
        persona_v2 = await generator_v2.generate_persona(
            template_data=template_data,
            mode=mode,
            gender=gender or None,
            custom_prompt=prompt or None
        )
        
        # Also generate V1 format for backward compatibility
        generator_v1 = EnhancedScenarioGenerator(azure_openai_client)
        generated_personas_v1 = await generator_v1.generate_personas_from_template(template_data, gender, prompt)
        
        # Return both formats
        personas = [{
            "v1": generated_personas_v1.get(persona_type, {}),
            "v2": persona_v2.dict()
        }] * count
        
        return {
            "template_id": template_id,
            "persona_type": persona_type,
            "mode": mode,
            "count": count,
            "personas": personas,
            "v2_categories_included": persona_v2.detail_categories_included
        }
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/generate-prompts-for-persona")
async def generate_prompts_for_persona(
    template_id: str = Body(...),
    persona_data: Dict[str, Any] = Body(...),
    modes: List[str] = Body(default=["assess", "try"]),
    db: Any = Depends(get_db)
):
    """
    Generate prompts for a specific persona.
    Call this for EACH persona you want to create prompts for.
    
    Args:
        template_id: Template to use
        persona_data: The persona object (from /generate-personas)
        modes: Which modes to generate (assess, try, learn)
    
    Returns:
        Prompts for the specified persona
    """
    try:
        from core.system_prompt_generator import SystemPromptGenerator
        from core.learn_mode_prompt_generator import LearnModePromptGenerator
        
        # Get template
        template = await db.templates.find_one({"id": template_id})
        if not template:
            raise HTTPException(status_code=404, detail="Template not found")
        
        template_data = template.get("template_data", {})
        
        # Initialize generators
        system_prompt_gen = SystemPromptGenerator(azure_openai_client, model="gpt-4o")
        learn_prompt_gen = LearnModePromptGenerator(azure_openai_client, model="gpt-4o")
        
        prompts = {}
        
        # Generate learn mode (no persona needed)
        if "learn" in modes:
            prompts["learn_mode"] = await learn_prompt_gen.generate_trainer_prompt(template_data)
        
        # Generate assess/try modes (persona-based, SAME prompt)
        if "assess" in modes or "try" in modes:
            # Generate ONE prompt for both assess and try
            system_prompt = await system_prompt_gen.generate_system_prompt(
                template_data=template_data,
                persona_data=persona_data,
                mode="assess_mode"  # Same for both
            )
            
            if "assess" in modes:
                prompts["assess_mode"] = system_prompt
            if "try" in modes:
                prompts["try_mode"] = system_prompt  # Same prompt
        
        return {
            "template_id": template_id,
            "persona_name": persona_data.get("name", "Unknown"),
            "prompts": prompts,
            "message": "Prompts generated for persona"
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/analyze-template-with-optional-docs")
async def analyze_template_with_optional_docs(
    template_file: UploadFile = File(...),
    supporting_docs: List[UploadFile] = File(default=[]),  # NOW OPTIONAL
    template_name: str = Form(...),
    current_user: UserDB = Depends(get_current_user),
    db: Any = Depends(get_db)
):
    """
    Updated version: Template file required, supporting docs optional
    """
    try:
        # 1. Process template file (same as before)
        content = await template_file.read()
        
        if template_file.filename.lower().endswith(('.doc', '.docx')):
            scenario_text = await extract_text_from_docx(content)
        elif template_file.filename.lower().endswith('.pdf'):
            scenario_text = await extract_text_from_pdf(content)
        else:
            scenario_text = content.decode('utf-8')
            
        # 2. Analyze template using V2 extractor
        from core.scenario_extractor_v2 import ScenarioExtractorV2
        extractor = ScenarioExtractorV2(azure_openai_client, model="gpt-4o")
        template_data = await extractor.extract_scenario_info(scenario_text)
        
        # 3. Process supporting documents ONLY if provided
        knowledge_base_id = None
        supporting_docs_metadata = []
        
        if supporting_docs and len(supporting_docs) > 0:
            knowledge_base_id = f"kb_{str(uuid4())}"
            supporting_docs_metadata = await process_and_index_documents(
                supporting_docs, knowledge_base_id, db
            )
        
        # 4. Save template (same structure for both flows)
        template_record = {
            "id": str(uuid4()),
            "name": template_name,
            "template_data": template_data,
            "knowledge_base_id": knowledge_base_id if supporting_docs_metadata else None,
            "supporting_documents": len(supporting_docs_metadata),
            "status": "ready_for_editing",
            "created_at": datetime.now().isoformat(),
            "updated_at": datetime.now().isoformat(),
            "created_by": str(current_user.id),
            "source": "file_upload"
        }
        
        await db.templates.insert_one(template_record)
        
        # 5. Create knowledge base if docs provided
        if supporting_docs_metadata:
            knowledge_base_record = {
                "_id": knowledge_base_id,
                "template_id": template_record["id"],
                "scenario_title": template_data.get("context_overview", {}).get("scenario_title", template_name),
                "supporting_documents": supporting_docs_metadata,
                "total_documents": len(supporting_docs_metadata),
                "total_chunks": sum(doc.get("chunk_count", 0) for doc in supporting_docs_metadata),
                "created_at": datetime.now(),
                "last_updated": datetime.now(),
                "fact_checking_enabled": True
            }
            await db.knowledge_bases.insert_one(knowledge_base_record)
        
        return {
            "template_id": template_record["id"],
            "template_data": template_data,
            "knowledge_base_id": knowledge_base_id,
            "supporting_documents_count": len(supporting_docs_metadata),
            "has_supporting_docs": len(supporting_docs_metadata) > 0,
            "fact_checking_enabled": len(supporting_docs_metadata) > 0,
            "message": "Template analyzed and ready for editing",
            "next_step": "Edit template sections, then generate prompts"
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error: {str(e)}")
@router.post("/analyze-scenario-enhanced")
async def analyze_scenario_enhanced(
    scenario_document: str = Form(...),
    template_name: str = Form(...),
    supporting_docs: List[UploadFile] = File(default=[]),  # OPTIONAL now
    current_user: UserDB = Depends(get_current_user),
    db: Any = Depends(get_db)
):
    """
    Enhanced analyze-scenario: Creates template + optionally processes documents
    - scenario_document: Text description of the scenario
    - template_name: Name for the template
    - supporting_docs: OPTIONAL supporting documents
    """
    try:
        # 1. Analyze scenario text using V2 extractor
        from core.scenario_extractor_v2 import ScenarioExtractorV2
        extractor = ScenarioExtractorV2(azure_openai_client, model="gpt-4o")
        template_data = await extractor.extract_scenario_info(scenario_document)
        
        # 2. Process supporting documents if provided (OPTIONAL)
        knowledge_base_id = None
        supporting_docs_metadata = []
        
        if supporting_docs and len(supporting_docs) > 0:
            knowledge_base_id = f"kb_{str(uuid4())}"
            supporting_docs_metadata = await process_and_index_documents(
                supporting_docs, knowledge_base_id, db
            )
        
        # 3. Create template record (same structure as analyze-template-with-docs)
        template_record = {
            "id": str(uuid4()),
            "name": template_name,
            "template_data": template_data,  # EDITABLE template data
            "knowledge_base_id": knowledge_base_id if supporting_docs_metadata else None,
            "supporting_documents": len(supporting_docs_metadata),
            "status": "ready_for_editing",
            "created_at": datetime.now().isoformat(),
            "updated_at": datetime.now().isoformat(),
            "created_by": str(current_user.id),
            "source": "text_analysis"  # vs "file_upload" for template files
        }
        
        await db.templates.insert_one(template_record)
        
        # 4. Create knowledge base record if documents were provided
        if supporting_docs_metadata:
            knowledge_base_record = {
                "_id": knowledge_base_id,
                "template_id": template_record["id"],
                "scenario_title": template_data.get("context_overview", {}).get("scenario_title", template_name),
                "supporting_documents": supporting_docs_metadata,
                "total_documents": len(supporting_docs_metadata),
                "total_chunks": sum(doc.get("chunk_count", 0) for doc in supporting_docs_metadata),
                "created_at": datetime.now(),
                "last_updated": datetime.now(),
                "fact_checking_enabled": True
            }
            await db.knowledge_bases.insert_one(knowledge_base_record)
        
        return {
            "template_id": template_record["id"],
            "template_data": template_data,  # For frontend editing
            "knowledge_base_id": knowledge_base_id,
            "supporting_documents_count": len(supporting_docs_metadata),
            "has_supporting_docs": len(supporting_docs_metadata) > 0,
            "fact_checking_enabled": len(supporting_docs_metadata) > 0,
            "message": "Scenario analyzed and template created successfully",
            "next_step": "Edit template sections, then generate prompts"
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error: {str(e)}")


@router.get("/scenarios/{scenario_id}/template-data")
async def get_scenario_template_data(
    scenario_id: str,
    current_user: UserDB = Depends(get_current_user),
    db: Any = Depends(get_db)
):
    """Get template data for a scenario"""
    try:
        # Get scenario
        scenario = await db.scenarios.find_one({"_id": scenario_id})
        if not scenario:
            raise HTTPException(status_code=404, detail="Scenario not found")
        
        # Get template from template_id
        template_id = scenario.get("template_id")
        if not template_id:
            raise HTTPException(status_code=404, detail="No template linked to this scenario")
        
        template = await db.templates.find_one({"id": template_id})
        if not template:
            raise HTTPException(status_code=404, detail="Template not found")
        
        template_data = template.get("template_data", {})
        knowledge_base_id = template.get("knowledge_base_id")
        
        return {
            "scenario_id": scenario_id,
            "template_id": template_id,
            "scenario_title": scenario.get("title", ""),
            "template_data": template_data,
            "knowledge_base_id": knowledge_base_id
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.put("/scenarios-editor/{scenario_id}/template-data")
async def update_scenario_template_data(
    scenario_id: str,
    request: Dict[str, Any] = Body(...),  # Accept raw dict instead of Pydantic model
    current_user: UserDB = Depends(get_current_user),
    db: Any = Depends(get_db)
):
    """Update template data for a scenario (updates the linked template)"""
    try:
        print('test')
        # Get scenario
        scenario = await db.scenarios.find_one({"_id": scenario_id})
        if not scenario:
            raise HTTPException(status_code=404, detail="Scenario not found")
        
        # Check permissions (similar to your existing scenario update logic)
        if current_user.role == UserRole.ADMIN:
            if scenario.get("created_by") != str(current_user.id):
                raise HTTPException(status_code=403, detail="Not authorized")
        elif current_user.role not in [UserRole.SUPERADMIN, UserRole.BOSS_ADMIN]:
            raise HTTPException(status_code=403, detail="Not authorized")
        
        # Get template_id
        template_id = scenario.get("template_id")
        if not template_id:
            raise HTTPException(status_code=404, detail="No template linked to this scenario")
        
        # Extract template_data from request (in case it's nested)
        template_data = request
        if "template_data" in request:
            template_data = request["template_data"]
        
        # Update the template document (not the scenario)
        await db.templates.update_one(
            {"id": template_id},
            {"$set": {
                "template_data": template_data,
                "updated_at": datetime.now().isoformat()
            }}
        )
        
        return {
            "message": "Template data updated successfully",
            "scenario_id": scenario_id,
            "template_id": template_id,
            "updated_template": template_data
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/scenarios/{scenario_id}/regenerate")
async def regenerate_scenario_from_template(
    scenario_id: str,
    regenerate_options: Dict[str, Any] = Body(default={
        "modes_to_regenerate": ["learn_mode", "assess_mode", "try_mode"],
        "regenerate_personas": True
    }),
    current_user: UserDB = Depends(get_current_user),
    db: Any = Depends(get_db)
):
    """Regenerate scenario prompts from linked template data"""
    try:
        # Get scenario
        scenario = await db.scenarios.find_one({"_id": scenario_id})
        if not scenario:
            raise HTTPException(status_code=404, detail="Scenario not found")
        
        # Check permissions
        if current_user.role == UserRole.ADMIN:
            if scenario.get("created_by") != str(current_user.id):
                raise HTTPException(status_code=403, detail="Not authorized")
        elif current_user.role not in [UserRole.SUPERADMIN, UserRole.BOSS_ADMIN]:
            raise HTTPException(status_code=403, detail="Not authorized")
        
        # Get template data from linked template
        template_id = scenario.get("template_id")
        if not template_id:
            raise HTTPException(status_code=400, detail="No template linked to this scenario")
        
        template = await db.templates.find_one({"id": template_id})
        if not template:
            raise HTTPException(status_code=404, detail="Linked template not found")
        
        # DEBUG: Log template fetch
        print(f"🔍 Regenerate: Fetched template {template_id}")
        print(f"📅 Template last updated: {template.get('updated_at')}")
        
        template_data = template.get("template_data")
        if not template_data:
            raise HTTPException(status_code=400, detail="No template data found")
        
        # DEBUG: Log template data preview
        print(f"📝 Template data keys: {list(template_data.keys())}")
        print(f"📋 Scenario title from template: {template_data.get('context_overview', {}).get('scenario_title', 'N/A')}")
        
        # Initialize generator
        generator = EnhancedScenarioGenerator(azure_openai_client)
        
        # Regenerate personas if requested
        personas = None
        if regenerate_options.get("regenerate_personas", True):
            personas = await generator.generate_personas_from_template(template_data)
        else:
            # Try to get existing personas (this might not exist in current schema)
            personas = scenario.get("generated_persona", {})
        
        # Generate prompts for requested modes
        modes_to_regenerate = regenerate_options.get("modes_to_regenerate", ["learn_mode", "assess_mode", "try_mode"])
        generated_prompts = {}
        
        if "learn_mode" in modes_to_regenerate:
            learn_prompt = await generator.generate_learn_mode_from_template(template_data)
            # learn_prompt = generator.insert_persona(learn_prompt, personas.get("learn_mode_expert", {}))
            # learn_prompt = generator.insert_language_instructions(learn_prompt, template_data.get("general_info", {}))
            generated_prompts["learn_mode"] = learn_prompt
        
        if "assess_mode" in modes_to_regenerate:
            assess_prompt = await generator.generate_assess_mode_from_template(template_data)
            # assess_prompt = generator.insert_persona(assess_prompt, personas.get("assess_mode_character", {}))
            # assess_prompt = generator.insert_language_instructions(assess_prompt, template_data.get("general_info", {}))
            generated_prompts["assess_mode"] = assess_prompt
        
        if "try_mode" in modes_to_regenerate:
            try_prompt = await generator.generate_try_mode_from_template(template_data)
            # try_prompt = generator.insert_persona(try_prompt, personas.get("assess_mode_character", {}))
            # try_prompt = generator.insert_language_instructions(try_prompt, template_data.get("general_info", {}))
            generated_prompts["try_mode"] = try_prompt
        

        
        update_data = {
            "updated_at": datetime.now()
        }
        
        # Store generated prompts in scenario AND update avatar_interaction documents
        print(f"🔍 DEBUG: Scenario modes to regenerate: {modes_to_regenerate}")
        print(f"🔍 DEBUG: Scenario learn_mode exists: {scenario.get('learn_mode') is not None}")
        print(f"🔍 DEBUG: Scenario assess_mode exists: {scenario.get('assess_mode') is not None}")
        print(f"🔍 DEBUG: Scenario try_mode exists: {scenario.get('try_mode') is not None}")
        
        if "learn_mode" in generated_prompts and scenario.get("learn_mode"):
            update_data["learn_mode.prompt"] = generated_prompts["learn_mode"]
            avatar_interaction_id = scenario.get("learn_mode", {}).get("avatar_interaction")
            print(f"🔍 DEBUG: Learn mode avatar_interaction: {avatar_interaction_id}")
            if not avatar_interaction_id:
                raise HTTPException(status_code=400, detail=f"Learn mode avatar_interaction is missing. Scenario learn_mode data: {scenario.get('learn_mode')}")
            try:
                result = await db.avatar_interactions.update_one(
                    {"_id": str(avatar_interaction_id)},
                    {"$set": {"system_prompt": generated_prompts["learn_mode"], "updated_at": datetime.now()}}
                )
                if result.matched_count == 0:
                    raise HTTPException(status_code=404, detail=f"Learn mode avatar_interaction {avatar_interaction_id} not found")
                print(f"✅ Learn mode avatar_interaction updated: matched={result.matched_count}, modified={result.modified_count}")
            except HTTPException:
                raise
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"Failed to update learn mode avatar_interaction: {str(e)}")
        
        if "assess_mode" in generated_prompts and scenario.get("assess_mode"):
            update_data["assess_mode.prompt"] = generated_prompts["assess_mode"]
            avatar_interaction_id = scenario.get("assess_mode", {}).get("avatar_interaction")
            print(f"🔍 DEBUG: Assess mode avatar_interaction: {avatar_interaction_id}")
            if not avatar_interaction_id:
                raise HTTPException(status_code=400, detail=f"Assess mode avatar_interaction is missing. Scenario assess_mode data: {scenario.get('assess_mode')}")
            try:
                result = await db.avatar_interactions.update_one(
                    {"_id": str(avatar_interaction_id)},
                    {"$set": {"system_prompt": generated_prompts["assess_mode"], "updated_at": datetime.now()}}
                )
                if result.matched_count == 0:
                    raise HTTPException(status_code=404, detail=f"Assess mode avatar_interaction {avatar_interaction_id} not found")
                print(f"✅ Assess mode avatar_interaction updated: matched={result.matched_count}, modified={result.modified_count}")
            except HTTPException:
                raise
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"Failed to update assess mode avatar_interaction: {str(e)}")
        
        if "try_mode" in generated_prompts and scenario.get("try_mode"):
            update_data["try_mode.prompt"] = generated_prompts["try_mode"]
            avatar_interaction_id = scenario.get("try_mode", {}).get("avatar_interaction")
            print(f"🔍 DEBUG: Try mode avatar_interaction: {avatar_interaction_id}")
            if not avatar_interaction_id:
                raise HTTPException(status_code=400, detail=f"Try mode avatar_interaction is missing. Scenario try_mode data: {scenario.get('try_mode')}")
            try:
                result = await db.avatar_interactions.update_one(
                    {"_id": str(avatar_interaction_id)},
                    {"$set": {"system_prompt": generated_prompts["try_mode"], "updated_at": datetime.now()}}
                )
                if result.matched_count == 0:
                    raise HTTPException(status_code=404, detail=f"Try mode avatar_interaction {avatar_interaction_id} not found")
                print(f"✅ Try mode avatar_interaction updated: matched={result.matched_count}, modified={result.modified_count}")
            except HTTPException:
                raise
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"Failed to update try mode avatar_interaction: {str(e)}")
        
        await db.scenarios.update_one(
            {"_id": scenario_id},
            {"$set": update_data}
        )
        
        return {
            "message": "Scenario regenerated successfully",
            "scenario_id": scenario_id,
            "template_id": template_id,
            "regenerated_modes": modes_to_regenerate,
            "generated_prompts": generated_prompts,
            "generated_personas": personas
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.put("/scenarios/{scenario_id}/knowledge-base")
async def manage_scenario_knowledge_base(
    scenario_id: str,
    kb_action: Dict[str, Any] = Body(...),
    current_user: UserDB = Depends(get_current_user),
    db: Any = Depends(get_db)
):
    """Link, unlink, or create knowledge base for scenario"""
    try:
        # Get scenario
        scenario = await db.scenarios.find_one({"_id": scenario_id})
        if not scenario:
            raise HTTPException(status_code=404, detail="Scenario not found")
        
        # Check permissions
        if current_user.role == UserRole.ADMIN:
            if scenario.get("created_by") != str(current_user.id):
                raise HTTPException(status_code=403, detail="Not authorized")
        elif current_user.role not in [UserRole.SUPERADMIN, UserRole.BOSS_ADMIN]:
            raise HTTPException(status_code=403, detail="Not authorized")
        
        action = kb_action.get("action")  # "link", "unlink", "create"
        knowledge_base_id = kb_action.get("knowledge_base_id")
        
        if action == "link":
            # Link existing knowledge base
            if not knowledge_base_id:
                raise HTTPException(status_code=400, detail="knowledge_base_id required for link action")
            
            # Verify KB exists
            kb = await db.knowledge_bases.find_one({"_id": knowledge_base_id})
            if not kb:
                raise HTTPException(status_code=404, detail="Knowledge base not found")
            
            await db.scenarios.update_one(
                {"_id": scenario_id},
                {"$set": {
                    "knowledge_base_id": knowledge_base_id,
                    "fact_checking_enabled": True,
                    "updated_at": datetime.now().isoformat()
                }}
            )
            
            return {
                "message": "Knowledge base linked successfully",
                "scenario_id": scenario_id,
                "knowledge_base_id": knowledge_base_id,
                "action": "linked"
            }
        
        elif action == "unlink":
            # Unlink knowledge base (but don't delete it)
            await db.scenarios.update_one(
                {"_id": scenario_id},
                {"$unset": {
                    "knowledge_base_id": "",
                    "fact_checking_enabled": ""
                },
                "$set": {
                    "updated_at": datetime.now().isoformat()
                }}
            )
            
            return {
                "message": "Knowledge base unlinked successfully",
                "scenario_id": scenario_id,
                "action": "unlinked"
            }
        
        elif action == "create":
            # Create new knowledge base for this scenario
            new_kb_id = f"kb_{scenario_id}_{str(uuid4())[:8]}"
            
            # Create empty knowledge base record
            kb_record = {
                "_id": new_kb_id,
                "scenario_id": scenario_id,
                "scenario_title": scenario.get("title", "Unknown Scenario"),
                "supporting_documents": [],
                "total_documents": 0,
                "created_at": datetime.now(),
                "last_updated": datetime.now(),
                "fact_checking_enabled": False  # Will be enabled when docs are added
            }
            
            await db.knowledge_bases.insert_one(kb_record)
            
            # Link to scenario
            await db.scenarios.update_one(
                {"_id": scenario_id},
                {"$set": {
                    "knowledge_base_id": new_kb_id,
                    "fact_checking_enabled": False,  # Will be enabled when docs are added
                    "updated_at": datetime.now().isoformat()
                }}
            )
            
            return {
                "message": "Knowledge base created and linked successfully",
                "scenario_id": scenario_id,
                "knowledge_base_id": new_kb_id,
                "action": "created"
            }
        
        else:
            raise HTTPException(status_code=400, detail="Invalid action. Use 'link', 'unlink', or 'create'")
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/scenarios/{scenario_id}/editing-interface")
async def get_scenario_editing_interface(
    scenario_id: str,
    current_user: UserDB = Depends(get_current_user),
    db: Any = Depends(get_db)
):
    """Get complete editing interface data for a scenario"""
    try:
        # Get scenario
        scenario = await db.scenarios.find_one({"_id": scenario_id})
        if not scenario:
            raise HTTPException(status_code=404, detail="Scenario not found")
        
        # Get template data from template_id
        template_id = scenario.get("template_id")
        template_data = {}
        
        if template_id:
            template = await db.templates.find_one({"id": template_id})
            if template:
                template_data = template.get("template_data", {})
        
        # Get knowledge base ID from template (not scenario)
        knowledge_base_id = None
        if template_id:
            template = await db.templates.find_one({"id": template_id})
            if template:
                knowledge_base_id = template.get("knowledge_base_id")
        
        # Get knowledge base info and documents
        kb_info = None
        kb_documents = []
        kb_stats = None
        
        if knowledge_base_id:
            # Get KB info
            kb_info = await db.knowledge_bases.find_one({"_id": knowledge_base_id})
            
            # Get KB documents
            docs_cursor = db.supporting_documents.find({"knowledge_base_id": knowledge_base_id})
            kb_documents = await docs_cursor.to_list(length=100)
            
            # Get KB stats
            total_chunks = sum(doc.get("chunk_count", 0) for doc in kb_documents)
            indexed_count = sum(1 for doc in kb_documents if doc.get("processing_status") == "completed")
            
            kb_stats = {
                "total_documents": len(kb_documents),
                "total_chunks": total_chunks,
                "indexed_documents": indexed_count,
                "search_ready": indexed_count > 0 and total_chunks > 0
            }
        
        return {
            "scenario_id": scenario_id,
            "scenario": {
                "title": scenario.get("title", ""),
                "description": scenario.get("description", ""),
                "template_id": template_id,
                "created_at": scenario.get("created_at"),
                "updated_at": scenario.get("updated_at")
            },
            "template_data": template_data,
            "knowledge_base": {
                "id": knowledge_base_id,
                "info": kb_info,
                "documents": kb_documents,
                "stats": kb_stats,
                "fact_checking_enabled": kb_info.get("fact_checking_enabled", False) if kb_info else False
            },
            "editing_capabilities": {
                "can_edit_template": True,
                "can_manage_knowledge_base": True,
                "can_regenerate": True,
                "available_modes": ["learn_mode", "assess_mode", "try_mode"]
            }
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    
 
# Add requirements.txt note for the Word document generation
"""
ADDITIONAL REQUIREMENTS FOR WORD GENERATION:
Add to your requirements.txt:

python-docx==0.8.11

For enhanced Word formatting, also consider:
python-docx2==0.1.0
docx2txt==0.8

Installation:
pip install python-docx==0.8.11
"""
# Add these imports at the top
from core.template_validator import TemplateValidator, PromptsValidator

# Add these validation endpoints



@router.post("/generate-prompts-from-template")
async def generate_prompts_from_template_with_validation(
    template_id: str = Body(...),
    modes: List[str] = Body(default=["learn", "try", "assess"]),
    validate_before_generation: bool = Body(default=True),
    current_user: UserDB = Depends(get_current_user),
    db: Any = Depends(get_db)
):
    """
    Generate prompts from template with optional validation
    """
    try:
        print(f"🔍 DEBUG: Starting generate_prompts_from_template with template_id: {template_id}")
        
        # Get template
        template = await db.templates.find_one({"id": template_id})
        if not template:
            print(f"❌ DEBUG: Template not found for id: {template_id}")
            raise HTTPException(status_code=404, detail="Template not found")
        
        print(f"✅ DEBUG: Template found, keys: {list(template.keys())}")
        
        # Get template_data and handle both string and dict formats
        template_data_raw = template.get("template_data", {})
        print(f"🔍 DEBUG: template_data_raw type: {type(template_data_raw)}")
        print(f"🔍 DEBUG: template_data_raw preview: {str(template_data_raw)[:200]}...")
        
        # Handle case where template_data is stored as JSON string
        if isinstance(template_data_raw, str):
            print(f"🔄 DEBUG: Converting string to dict")
            try:
                import json
                template_data = json.loads(template_data_raw)
                print(f"✅ DEBUG: Successfully parsed JSON, type: {type(template_data)}")
            except json.JSONDecodeError as e:
                print(f"❌ DEBUG: JSON decode error: {e}")
                raise HTTPException(status_code=400, detail="Invalid template data format - not valid JSON")
        elif isinstance(template_data_raw, dict):
            print(f"✅ DEBUG: template_data_raw is already a dict")
            template_data = template_data_raw
        else:
            print(f"❌ DEBUG: Unexpected template_data type: {type(template_data_raw)}")
            raise HTTPException(status_code=400, detail=f"Template data must be a dictionary or JSON string, got {type(template_data_raw)}")
        
        # Final validation that we have a dictionary
        if not isinstance(template_data, dict):
            print(f"❌ DEBUG: Final template_data is not a dict: {type(template_data)}")
            raise HTTPException(status_code=400, detail="Template data must be a dictionary after parsing")
        
        print(f"✅ DEBUG: Final template_data is dict with keys: {list(template_data.keys())}")
        
        # Validate template if requested
        validation_result = None
        if validate_before_generation:
            print(f"🔍 DEBUG: Starting template validation")
            try:
                validation_result = TemplateValidator.validate_template(template_data)
                print(f"✅ DEBUG: Template validation completed, valid: {validation_result.valid}")
            except Exception as e:
                print(f"❌ DEBUG: Template validation error: {e}")
                print(f"❌ DEBUG: Error type: {type(e)}")
                import traceback
                print(f"❌ DEBUG: Traceback: {traceback.format_exc()}")
                raise HTTPException(status_code=500, detail=f"Template validation failed: {str(e)}")
            
            if not validation_result.valid:
                print(f"⚠️ DEBUG: Template validation failed")
                return {
                    "error": "Template validation failed",
                    "validation": {
                        "valid": False,
                        "score": validation_result.score,
                        "issues": [
                            {
                                "field": issue.field,
                                "severity": issue.severity,
                                "message": issue.message,
                                "suggestion": issue.suggestion
                            }
                            for issue in validation_result.issues
                        ]
                    },
                    "message": "Fix template errors before generating prompts"
                }
        
        # Generate prompts
        generator = EnhancedScenarioGenerator(azure_openai_client)
        
        # Get archetype classification
        archetype_classification = template_data.get("archetype_classification")
        
        # Generate personas
        personas = await generator.generate_personas_from_template(
            template_data, "", "", archetype_classification
        )
        
        # Generate prompts for requested modes
        prompts = {}
        if "learn" in modes:
            prompts["learn_mode_prompt"] = await generator.generate_learn_mode_from_template(template_data)
        if "try" in modes:
            prompts["try_mode_prompt"] = await generator.generate_try_mode_from_template(template_data)
        if "assess" in modes:
            prompts["assess_mode_prompt"] = await generator.generate_assess_mode_from_template(template_data)
        
        prompts["personas"] = personas
        
        return {
            "template_id": template_id,
            "prompts": prompts,
            "message": "Prompts generated successfully"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error: {str(e)}")




@router.put("/templates/{template_id}")
async def update_template(
    template_id: str,
    template_data: Dict[str, Any] = Body(...),
    db: Any = Depends(get_db)
):
    """Update template - matches frontend expectation"""
    try:
        result = await db.templates.update_one(
            {"id": template_id},
            {"$set": {
                "template_data": template_data,
                "updated_at": datetime.now().isoformat()
            }}
        )
        
        if result.matched_count == 0:
            raise HTTPException(status_code=404, detail="Template not found")
        
        return {
            "message": "Template updated successfully",
            "template_id": template_id
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


"""
Add these endpoints to scenario_generator.py
"""

from fastapi import APIRouter, Depends, HTTPException, Body
from typing import Any, List, Optional
from uuid import UUID
from core.persona_db_manager import PersonaDBManager
from models.persona_models import PersonaV2Response, PersonaInstanceV2

# Add to your existing router in scenario_generator.py


@router.post("/personas/v2/save")
async def save_persona_v2(
    request_body: dict = Body(...),
    current_user: UserDB = Depends(get_current_user),
    db: Any = Depends(get_db)
):
    """
    Save persona with ANY fields dynamically - no Pydantic validation.
    Just save whatever fields are sent.
    """
    try:
        # Extract persona_data from request body
        persona_data = request_body.get("persona_data", request_body)
        
        # Add metadata fields
        persona_data["id"] = persona_data.get("id", str(uuid4()))
        persona_data["created_by"] = str(current_user.id)
        persona_data["created_at"] = datetime.now().isoformat()
        persona_data["updated_at"] = datetime.now().isoformat()
        
        # Save directly to database - no validation
        await db.personas.insert_one(persona_data)
        
        return {
            "success": True,
            "persona_id": persona_data["id"],
            "message": f"Saved persona: {persona_data.get('name', 'Unknown')}",
            "fields_saved": list(persona_data.keys())
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.put("/personas/v2/update/{persona_id}")
async def update_persona_v2(
    persona_id: str,
    request_body: dict = Body(...),
    current_user: UserDB = Depends(get_current_user),
    db: Any = Depends(get_db)
):
    """
    Update existing persona with ANY fields dynamically.
    """
    try:
        persona_data = request_body.get("persona_data", request_body)
        persona_data["updated_at"] = datetime.now().isoformat()
        
        # Update in personas_v2 collection (use _id)
        result = await db.personas_v2.update_one(
            {"_id": persona_id},
            {"$set": persona_data}
        )
        
        # Fallback: try personas collection (use id)
        if result.matched_count == 0:
            result = await db.personas.update_one(
                {"id": persona_id},
                {"$set": persona_data}
            )
        
        if result.matched_count == 0:
            raise HTTPException(status_code=404, detail="Persona not found")
        
        return {
            "success": True,
            "persona_id": persona_id,
            "message": f"Updated persona: {persona_data.get('name', 'Unknown')}"
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/personas/v2/{persona_id}")
async def get_persona_v2(
    persona_id: str,
    db: Any = Depends(get_db)
):
    """Get persona by ID - returns raw data"""
    try:
        persona = await db.personas.find_one({"id": persona_id})
        
        if not persona:
            raise HTTPException(status_code=404, detail="Persona not found")
        
        # Convert ObjectId to string
        if "_id" in persona:
            persona["_id"] = str(persona["_id"])
        
        return persona
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/personas/v2/list")
async def list_personas_v2(
    template_id: Optional[str] = None,
    mode: Optional[str] = None,
    archetype: Optional[str] = None,
    limit: int = 50,
    db: Any = Depends(get_db)
):
    """
    List personas with optional filters - returns raw data.
    """
    try:
        # Build filter query
        query = {}
        if template_id:
            query["template_id"] = template_id
        if mode:
            query["mode"] = mode
        if archetype:
            query["archetype"] = archetype
        
        # Get personas from database
        cursor = db.personas.find(query).limit(limit)
        personas = await cursor.to_list(length=limit)
        
        # Convert ObjectId to string
        for persona in personas:
            if "_id" in persona:
                persona["_id"] = str(persona["_id"])
        
        return personas
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/personas/v2/template/{template_id}")
async def get_personas_by_template(
    template_id: str,
    db: Any = Depends(get_db)
):
    """Get all personas for a specific template"""
    try:
        personas = await PersonaDBManager.get_personas_by_template(db, template_id)
        
        return {
            "success": True,
            "template_id": template_id,
            "count": len(personas),
            "personas": [p.dict() for p in personas]
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/personas/v2/search/category/{category_name}")
async def search_by_category(
    category_name: str,
    limit: int = 20,
    db: Any = Depends(get_db)
):
    """
    Search personas that have a specific detail category.
    
    Example: /personas/v2/search/category/sales_rep_history
    Returns all personas with sales_rep_history category.
    """
    try:
        personas = await PersonaDBManager.search_personas_by_category(
            db, category_name, limit
        )
        
        return {
            "success": True,
            "category": category_name,
            "count": len(personas),
            "personas": personas
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.put("/personas/v2/{persona_id}")
async def update_persona_v2(
    persona_id: str,
    updates: dict = Body(...),
    db: Any = Depends(get_db)
):
    """
    Update persona fields - accepts ANY fields dynamically.
    """
    try:
        # Add updated_at timestamp
        updates["updated_at"] = datetime.now().isoformat()
        
        # Update in database
        result = await db.personas.update_one(
            {"id": persona_id},
            {"$set": updates}
        )
        
        if result.matched_count == 0:
            raise HTTPException(status_code=404, detail="Persona not found")
        
        return {
            "success": True,
            "message": "Persona updated successfully",
            "fields_updated": list(updates.keys())
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/personas/v2/{persona_id}")
async def delete_persona_v2(
    persona_id: str,
    db: Any = Depends(get_db)
):
    """Delete persona"""
    try:
        result = await db.personas.delete_one({"id": persona_id})
        
        if result.deleted_count == 0:
            raise HTTPException(status_code=404, detail="Persona not found")
        
        return {
            "success": True,
            "message": "Persona deleted successfully"
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/personas/v2/generate-and-save")
async def generate_and_save_persona(
    template_id: str = Body(...),
    persona_type_index: int = Body(default=0),
    gender: Optional[str] = Body(default=None),
    custom_prompt: Optional[str] = Body(default=None),
    current_user: UserDB = Depends(get_current_user),
    db: Any = Depends(get_db)
):
    """
    Generate 1 persona for try/assess mode (shared).
    Learn mode doesn't need persona.
    
    Steps:
    1. Load template_data from database
    2. Generate persona using PersonaGeneratorV2 (try_mode)
    3. Save to database with mode="try_assess_mode"
    4. Return persona with ID
    """
    try:
        from core.persona_generator_v2 import PersonaGeneratorV2
        
        # Load template
        template_doc = await db.templates.find_one({"id": template_id})
        if not template_doc:
            raise HTTPException(status_code=404, detail="Template not found")
        
        template_data = template_doc.get("template_data", {})
        template_data["template_id"] = template_id
        
        # Get persona type from template
        persona_types = template_data.get("persona_types", [])
        if not persona_types or persona_type_index >= len(persona_types):
            raise HTTPException(status_code=400, detail="Invalid persona_type_index or no persona types defined")
        
        selected_persona_type = persona_types[persona_type_index]
        
        # Generate persona for try/assess (shared)
        generator = PersonaGeneratorV2(azure_openai_client)
        persona = await generator.generate_persona(
            template_data=template_data,
            mode="assess_mode",
            gender=gender,
            custom_prompt=custom_prompt
        )
        print(persona,"modeeeee")
        # Override description and role, but keep the generated realistic name
        persona.description = selected_persona_type.get("description", "")
        persona.role = selected_persona_type.get("type", "Character")
        # persona.name is already set by PersonaGeneratorV2 to a realistic name
        
        # Save to database
        persona_id = await PersonaDBManager.save_persona(db, persona, current_user.id)
        
        return {
            "success": True,
            "persona_id": persona_id,
            "persona": persona.dict(),
            "message": f"Generated persona: {persona.name}"
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
