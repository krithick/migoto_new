"""
Phase 1: Basic Document Processing
Handles text extraction, chunking, and basic vector preparation
"""

import asyncio
import io
import os
from typing import List, Dict, Any, Optional, Tuple
from uuid import uuid4
from datetime import datetime
import json
import re
from core.simple_token_logger import log_embedding_usage,log_token_usage

# Text extraction imports
try:
    import pypdf
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# OpenAI for embeddings
from openai import AsyncAzureOpenAI

# Import models
from models.enhanced_models import (
    DocumentChunk, DocumentMetadata, ProcessingStatus, 
    prepare_for_mongodb, Collections
)

class DocumentProcessor:
    """Handles document text extraction and chunking"""
    
    def __init__(self, openai_client: AsyncAzureOpenAI, db: Any):
        self.openai_client = openai_client
        self.db = db
        self.chunk_size = 1000  # Characters per chunk
        self.chunk_overlap = 200  # Overlap between chunks
        
    async def process_document(self, document_id: str, file_content: bytes, 
                             filename: str) -> List[DocumentChunk]:
        """Process document: extract text, create chunks, generate embeddings"""
        
        # Extract text based on file type
        text_content = await self._extract_text(file_content, filename)
        
        if not text_content or len(text_content.strip()) < 50:
            raise ValueError(f"Insufficient content extracted from {filename}")
        
        # Create chunks
        chunks = await self._create_chunks(text_content, document_id, filename)
        
        # Generate embeddings for chunks
        chunks_with_embeddings = await self._generate_embeddings(chunks)
        
        return chunks_with_embeddings
    
    async def _extract_text(self, file_content: bytes, filename: str) -> str:
        """Extract text from various file formats"""
        
        file_ext = os.path.splitext(filename)[1].lower()
        
        if file_ext == '.pdf':
            return await self._extract_pdf_text(file_content)
        elif file_ext in ['.docx', '.doc']:
            return await self._extract_docx_text(file_content)
        elif file_ext == '.txt':
            return file_content.decode('utf-8', errors='ignore')
        else:
            # Try to decode as text
            try:
                return file_content.decode('utf-8')
            except UnicodeDecodeError:
                return file_content.decode('utf-8', errors='ignore')
    
    async def _extract_pdf_text(self, file_content: bytes) -> str:
        """Extract text from PDF files"""
        if not HAS_PYPDF:
            raise ImportError("pypdf package is required for PDF processing")
        
        try:
            with io.BytesIO(file_content) as pdf_file:
                reader = pypdf.PdfReader(pdf_file)
                text_parts = []
                
                for page_num, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text.strip():
                        # Add page number for reference
                        text_parts.append(f"[Page {page_num + 1}]\n{page_text}")
                
                return "\n\n".join(text_parts)
                
        except Exception as e:
            raise ValueError(f"Error extracting PDF text: {str(e)}")
    
    async def _extract_docx_text(self, file_content: bytes) -> str:
        """Extract text from DOCX files"""
        if not HAS_DOCX:
            raise ImportError("python-docx package is required for DOCX processing")
        
        try:
            with io.BytesIO(file_content) as docx_file:
                doc = docx.Document(docx_file)
                text_parts = []
                
                # Extract paragraphs
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text.strip())
                
                # Extract tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            text_parts.append(" | ".join(row_text))
                
                return "\n\n".join(text_parts)
                
        except Exception as e:
            raise ValueError(f"Error extracting DOCX text: {str(e)}")
    
    async def _create_chunks(self, text: str, document_id: str, 
                           filename: str) -> List[DocumentChunk]:
        """Create overlapping chunks from text"""
        
        # Clean and normalize text
        text = self._clean_text(text)
        
        # Split into sentences for better chunking
        sentences = self._split_into_sentences(text)
        
        chunks = []
        current_chunk = ""
        current_word_count = 0
        chunk_index = 0
        
        for sentence in sentences:
            sentence_length = len(sentence)
            sentence_words = len(sentence.split())
            
            # Check if adding this sentence would exceed chunk size
            if (len(current_chunk) + sentence_length > self.chunk_size and 
                current_chunk.strip()):
                
                # Create chunk
                chunk = DocumentChunk(
                    document_id=document_id,
                    knowledge_base_id="temp",  # Will be set by caller
                    content=current_chunk.strip(),
                    chunk_index=chunk_index,
                    word_count=current_word_count,
                    character_count=len(current_chunk.strip()),
                    metadata={
                        "source_file": filename,
                        "extraction_method": "sentence_based"
                    }
                )
                chunks.append(chunk)
                
                # Start new chunk with overlap
                overlap_text = self._get_overlap_text(current_chunk, self.chunk_overlap)
                current_chunk = overlap_text + " " + sentence
                current_word_count = len(current_chunk.split())
                chunk_index += 1
            else:
                current_chunk += " " + sentence if current_chunk else sentence
                current_word_count += sentence_words
        
        # Add final chunk
        if current_chunk.strip():
            chunk = DocumentChunk(
                document_id=document_id,
                knowledge_base_id="temp",  # Will be set by caller
                content=current_chunk.strip(),
                chunk_index=chunk_index,
                word_count=current_word_count,
                character_count=len(current_chunk.strip()),
                metadata={
                    "source_file": filename,
                    "extraction_method": "sentence_based"
                }
            )
            chunks.append(chunk)
        
        return chunks
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove page headers/footers (common patterns)
        text = re.sub(r'\[Page \d+\]\s*', '', text)
        
        # Remove excessive line breaks
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        return text.strip()
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences"""
        # Simple sentence splitting (can be improved with NLTK)
        sentences = re.split(r'(?<=[.!?])\s+', text)
        
        # Filter out very short sentences
        sentences = [s.strip() for s in sentences if len(s.strip()) > 10]
        
        return sentences
    
    def _get_overlap_text(self, text: str, overlap_size: int) -> str:
        """Get overlap text from end of chunk"""
        if len(text) <= overlap_size:
            return text
        
        # Get last overlap_size characters, but try to break at word boundary
        overlap_text = text[-overlap_size:]
        
        # Find last complete word
        last_space = overlap_text.rfind(' ')
        if last_space > overlap_size // 2:  # If we have a reasonable word break
            overlap_text = overlap_text[last_space + 1:]
        
        return overlap_text
    
    async def _generate_embeddings(self, chunks: List[DocumentChunk]) -> List[DocumentChunk]:
        """Generate embeddings for document chunks"""
        
        if not chunks:
            return chunks
        
        # Prepare texts for embedding
        texts = [chunk.content for chunk in chunks]
        
        try:
            # Generate embeddings in batches
            batch_size = 100  # Azure OpenAI limit
            all_embeddings = []
            
            for i in range(0, len(texts), batch_size):
                batch_texts = texts[i:i + batch_size]
                batch_embeddings = await self._get_embeddings_batch(batch_texts)
                all_embeddings.extend(batch_embeddings)
            
            # Assign embeddings to chunks
            for chunk, embedding in zip(chunks, all_embeddings):
                chunk.embedding = embedding
            
            return chunks
            
        except Exception as e:
            # Log error but don't fail - embeddings can be generated later
            print(f"Warning: Failed to generate embeddings: {e}")
            return chunks
    
    async def _get_embeddings_batch(self, texts: List[str]) -> List[List[float]]:
        """Get embeddings for batch of texts"""
        
        response = await self.openai_client.embeddings.create(
            input=texts,
            model="text-embedding-ada-002"
        )
        log_embedding_usage(response, "get_embeddings_batch", len(texts))
        return [data.embedding for data in response.data]

class KnowledgeBaseManager:
    """Manages knowledge base creation and updates"""
    
    def __init__(self, db: Any, document_processor: DocumentProcessor):
        self.db = db
        self.document_processor = document_processor
    
    async def add_document_to_knowledge_base(self, knowledge_base_id: str, 
                                          document_id: str, file_content: bytes, 
                                          filename: str) -> Dict[str, Any]:
        """Add document to knowledge base"""
        
        try:
            # Process document
            chunks = await self.document_processor.process_document(
                document_id, file_content, filename
            )
            
            # Set knowledge base ID for all chunks
            for chunk in chunks:
                chunk.knowledge_base_id = knowledge_base_id
            
            # Store chunks in database
            chunk_docs = [prepare_for_mongodb(chunk) for chunk in chunks]
            
            # Use a custom collection for chunks (not in your existing collections)
            if chunk_docs:
                await self.db.document_chunks.insert_many(chunk_docs)
            
            # Update document metadata
            await self.db[Collections.DOCUMENT_METADATA].update_one(
                {"_id": document_id},
                {"$set": {
                    "chunk_count": len(chunks),
                    "chunk_ids": [chunk.id for chunk in chunks],
                    "processing_status": ProcessingStatus.COMPLETED.value,
                    "processed_at": datetime.now()
                }}
            )
            
            # Update knowledge base
            await self.db[Collections.KNOWLEDGE_BASES].update_one(
                {"_id": knowledge_base_id},
                {
                    "$inc": {"total_chunks": len(chunks)},
                    "$set": {"last_updated": datetime.now()}
                }
            )
            
            return {
                "success": True,
                "chunks_created": len(chunks),
                "chunks_with_embeddings": len([c for c in chunks if c.embedding]),
                "document_id": document_id
            }
            
        except Exception as e:
            # Update document with error status
            await self.db[Collections.DOCUMENT_METADATA].update_one(
                {"_id": document_id},
                {"$set": {
                    "processing_status": ProcessingStatus.FAILED.value,
                    "error_message": str(e)
                }}
            )
            
            raise e
    
    async def search_knowledge_base(self, knowledge_base_id: str, query: str, 
                                  top_k: int = 5) -> List[Dict[str, Any]]:
        """Search knowledge base (basic text search for now)"""
        
        try:
            # For now, do basic text search on chunks
            # In production, this would use Azure Search vector search
            
            pipeline = [
                {"$match": {"knowledge_base_id": knowledge_base_id}},
                {"$addFields": {
                    "relevance_score": {
                        "$sum": [
                            {"$cond": [
                                {"$regexMatch": {
                                    "input": "$content",
                                    "regex": re.escape(query),
                                    "options": "i"
                                }}, 
                                10, 0
                            ]},
                            {"$multiply": [
                                {"$size": {
                                    "$filter": {
                                        "input": {"$split": ["$content", " "]},
                                        "cond": {
                                            "$regexMatch": {
                                                "input": "$this",
                                                "regex": re.escape(query),
                                                "options": "i"
                                            }
                                        }
                                    }
                                }},
                                5
                            ]}
                        ]
                    }
                }},
                {"$match": {"relevance_score": {"$gt": 0}}},
                {"$sort": {"relevance_score": -1}},
                {"$limit": top_k}
            ]
            
            results = await self.db.document_chunks.aggregate(pipeline).to_list(length=top_k)
            
            return [
                {
                    "content": result["content"],
                    "chunk_id": result["_id"],
                    "document_id": result["document_id"],
                    "relevance_score": result["relevance_score"],
                    "metadata": result.get("metadata", {})
                }
                for result in results
            ]
            
        except Exception as e:
            print(f"Error searching knowledge base: {e}")
            return []
    
    async def get_knowledge_base_stats(self, knowledge_base_id: str) -> Dict[str, Any]:
        """Get statistics about knowledge base"""
        
        try:
            # Get knowledge base info
            kb_doc = await self.db[Collections.KNOWLEDGE_BASES].find_one(
                {"_id": knowledge_base_id}
            )
            
            if not kb_doc:
                return {"error": "Knowledge base not found"}
            
            # Get document count
            doc_count = await self.db[Collections.DOCUMENT_METADATA].count_documents(
                {"knowledge_base_id": knowledge_base_id}
            )
            
            # Get chunk count
            chunk_count = await self.db.document_chunks.count_documents(
                {"knowledge_base_id": knowledge_base_id}
            )
            
            # Get processing status distribution
            status_pipeline = [
                {"$match": {"knowledge_base_id": knowledge_base_id}},
                {"$group": {
                    "_id": "$processing_status",
                    "count": {"$sum": 1}
                }}
            ]
            
            status_results = await self.db[Collections.DOCUMENT_METADATA].aggregate(
                status_pipeline
            ).to_list(length=10)
            
            status_distribution = {result["_id"]: result["count"] for result in status_results}
            
            return {
                "knowledge_base_id": knowledge_base_id,
                "total_documents": doc_count,
                "total_chunks": chunk_count,
                "processing_status": status_distribution,
                "last_updated": kb_doc.get("last_updated"),
                "created_at": kb_doc.get("created_at")
            }
            
        except Exception as e:
            return {"error": str(e)}

class BasicFactChecker:
    """Basic fact-checking functionality for Phase 1"""
    
    def __init__(self, knowledge_manager: KnowledgeBaseManager, openai_client: AsyncAzureOpenAI):
        self.knowledge_manager = knowledge_manager
        self.openai_client = openai_client
    
    async def extract_claims_from_text(self, text: str) -> List[str]:
        """Extract factual claims from text using LLM"""
        
        try:
            prompt = f"""
            Extract specific factual claims from the following text that can be verified against a knowledge base.
            Focus on concrete statements about:
            - Prices, costs, or financial information
            - Product features or specifications  
            - Policies or procedures
            - Dates, times, or schedules
            - Contact information
            - Technical details
            
            Text: {text}
            
            Return only the factual claims as a JSON array of strings. Ignore opinions, suggestions, or general statements.
            """
            
            response = await self.openai_client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "You extract factual claims that can be verified."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.1,
                max_tokens=500
            )
            log_token_usage(response,"extract_claims_from_text")
            result_text = response.choices[0].message.content
            
            # Try to parse JSON
            try:
                claims = json.loads(result_text)
                return claims if isinstance(claims, list) else []
            except json.JSONDecodeError:
                # Fallback: extract lines that look like claims
                lines = result_text.split('\n')
                claims = []
                for line in lines:
                    line = line.strip()
                    if line and not line.startswith('-') and len(line) > 10:
                        claims.append(line)
                return claims
                
        except Exception as e:
            print(f"Error extracting claims udjk: {e}")
            return []
    
    async def verify_claim(self, claim: str, knowledge_base_id: str) -> Dict[str, Any]:
        """Verify a claim against knowledge base"""
        
        try:
            # Search for relevant information
            search_results = await self.knowledge_manager.search_knowledge_base(
                knowledge_base_id, claim, top_k=3
            )
            
            if not search_results:
                return {
                    "claim": claim,
                    "result": "UNSUPPORTED",
                    "confidence": 0.0,
                    "explanation": "No relevant information found in knowledge base",
                    "supporting_evidence": []
                }
            
            # Use LLM to verify claim against search results
            context = "\n\n".join([
                f"Source: {result['metadata'].get('source_file', 'Unknown')}\n{result['content']}"
                for result in search_results
            ])
            
            verification_prompt = f"""
            Verify the following claim against the provided knowledge base information:
            
            CLAIM: {claim}
            
            KNOWLEDGE BASE:
            {context}
            
            Analyze and provide your assessment as JSON:
            {{
                "result": "CORRECT|INCORRECT|PARTIALLY_CORRECT|UNSUPPORTED|UNCLEAR",
                "confidence": 0.0-1.0,
                "explanation": "Detailed explanation of your assessment",
                "supporting_evidence": ["Evidence from knowledge base that supports or contradicts the claim"]
            }}
            
            Guidelines:
            - CORRECT: Claim is fully accurate according to the knowledge base
            - INCORRECT: Claim contains factual errors
            - PARTIALLY_CORRECT: Claim has some accurate elements but also errors
            - UNSUPPORTED: Knowledge base doesn't contain enough information
            - UNCLEAR: Claim is too vague to verify
            """
            
            response = await self.openai_client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "You are a precise fact-checking assistant."},
                    {"role": "user", "content": verification_prompt}
                ],
                temperature=0.1,
                max_tokens=800
            )
            log_token_usage(response,"verify_claim")
            result_text = response.choices[0].message.content
            
            try:
                verification_result = json.loads(result_text)
                verification_result["claim"] = claim
                verification_result["search_results"] = search_results
                return verification_result
            except json.JSONDecodeError:
                # Fallback result
                return {
                    "claim": claim,
                    "result": "UNCLEAR",
                    "confidence": 0.5,
                    "explanation": "Could not properly analyze the claim",
                    "supporting_evidence": [],
                    "search_results": search_results
                }
                
        except Exception as e:
            return {
                "claim": claim,
                "result": "ERROR",
                "confidence": 0.0,
                "explanation": f"Error during verification: {str(e)}",
                "supporting_evidence": []
            }

# Utility functions
async def setup_document_processing():
    """Setup document processing components"""
    
    # Check required packages
    missing_packages = []
    if not HAS_PYPDF:
        missing_packages.append("pypdf")
    if not HAS_DOCX:
        missing_packages.append("python-docx")
    
    if missing_packages:
        print(f"Warning: Missing packages for document processing: {missing_packages}")
        print("Install with: pip install " + " ".join(missing_packages))
    
    # Create necessary directories
    os.makedirs("uploads/enhanced_scenarios", exist_ok=True)
    
    print("✅ Document processing setup completed")

# Enhanced background processing function
async def enhanced_process_document_background(document_id: str, job_id: str, 
                                            file_content: bytes, filename: str,
                                            knowledge_base_id: str, job_manager, db):
    """Enhanced background document processing with real functionality"""
    
    try:
        from openai import AsyncAzureOpenAI
        
        # Initialize components
        openai_client = AsyncAzureOpenAI(
            api_key=os.getenv("api_key"),
            azure_endpoint=os.getenv("endpoint"),
            api_version=os.getenv("api_version")
        )
        
        document_processor = DocumentProcessor(openai_client, db)
        knowledge_manager = KnowledgeBaseManager(db, document_processor)
        
        await job_manager.update_job_progress(
            job_id, 10, "Initializing document processing...", ProcessingStatus.PROCESSING
        )
        
        await job_manager.update_job_progress(job_id, 30, "Extracting text from document...")
        
        # Process document with real functionality
        result = await knowledge_manager.add_document_to_knowledge_base(
            knowledge_base_id, document_id, file_content, filename
        )
        
        await job_manager.update_job_progress(job_id, 60, "Creating document chunks...")
        await asyncio.sleep(0.5)  # Brief pause for progress update
        
        await job_manager.update_job_progress(job_id, 80, "Generating embeddings...")
        await asyncio.sleep(0.5)
        
        await job_manager.update_job_progress(job_id, 95, "Finalizing...")
        
        # Update job with results
        await job_manager.update_job_progress(
            job_id, 100, 
            f"Document processing completed. Created {result['chunks_created']} chunks.",
            ProcessingStatus.COMPLETED
        )
        
        # Store job result
        await db[Collections.PROCESSING_JOBS].update_one(
            {"_id": job_id},
            {"$set": {"result": result}}
        )
        
    except Exception as e:
        await job_manager.mark_job_failed(job_id, str(e))
        print(f"Document processing failed: {e}")

# Test functions
async def test_document_processing():
    """Test document processing functionality"""
    
    print("Testing document processing...")
    
    # Test with sample text file
    sample_text = """
    Premium Maternity Package
    
    Our premium maternity package is designed for expecting families who want the highest level of care and comfort.
    
    Package Details:
    - Price Range: ₹85,000 - ₹1,20,000
    - Private AC room with attached bathroom
    - Dedicated nursing care during day shift (8 AM - 8 PM)
    - Daily consultations with OB-GYN specialist
    - Pediatrician consultation (3 sessions included)
    - Complimentary photography session (2 hours)
    - All meals for patient and one attendant
    
    Exclusions:
    - ICU charges if complications arise
    - Additional specialist consultations beyond included sessions
    - Extended photography sessions
    
    Insurance Information:
    - Cashless facility available with major insurance providers
    - Pre-authorization required 48 hours in advance
    - Coverage varies by insurance policy
    """
    
    try:
        from openai import AsyncAzureOpenAI
        
        # Mock OpenAI client for testing
        class MockOpenAIClient:
            async def embeddings_create(self, **kwargs):
                # Return mock embeddings
                return type('Response', (), {
                    'data': [type('EmbeddingData', (), {'embedding': [0.1] * 1536})() 
                           for _ in kwargs['input']]
                })()
        
        mock_db = type('MockDB', (), {
            'document_chunks': type('Collection', (), {
                'insert_many': lambda docs: None
            })()
        })()
        
        processor = DocumentProcessor(MockOpenAIClient(), mock_db)
        
        # Test text extraction
        chunks = await processor._create_chunks(sample_text, "test_doc", "test.txt")
        
        print(f"✅ Created {len(chunks)} chunks from sample text")
        print(f"✅ First chunk: {chunks[0].content[:100]}...")
        print(f"✅ Chunk word counts: {[chunk.word_count for chunk in chunks]}")
        
        return True
        
    except Exception as e:
        print(f"❌ Test failed: {e}")
        return False

if __name__ == "__main__":
    # Run tests
    asyncio.run(test_document_processing())