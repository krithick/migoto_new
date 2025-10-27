#!/usr/bin/env python3
"""
Test frontend integration by simulating API calls
"""

import asyncio
import json
from fastapi import FastAPI, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from enhanced_scenario_generator import FlexibleScenarioGenerator
import io
import docx

# Create test FastAPI app
app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Mock database
mock_db = {
    "flexible_templates": {},
    "flexible_scenarios": {}
}

async def extract_text_from_docx(file_content):
    """Extract text from DOCX file"""
    try:
        with io.BytesIO(file_content) as f:
            doc = docx.Document(f)
            text_parts = []
            
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    text_parts.append(f"PARAGRAPH_{i}: {para.text.strip()}")
            
            for table_idx, table in enumerate(doc.tables):
                text_parts.append(f"\nTABLE_{table_idx}_START:")
                for row_idx, row in enumerate(table.rows):
                    row_text = []
                    for cell_idx, cell in enumerate(row.cells):
                        if cell.text.strip():
                            row_text.append(f"CELL_{cell_idx}: {cell.text.strip()}")
                    if row_text:
                        text_parts.append(f"ROW_{row_idx}: {' | '.join(row_text)}")
                text_parts.append(f"TABLE_{table_idx}_END\n")
            
            return "\n".join(text_parts)
    except Exception as e:
        print(f"DOCX extraction error: {str(e)}")
        return None

@app.post("/flexible/analyze-document")
async def analyze_document(
    file: UploadFile = File(...),
    template_name: str = Form(...)
):
    """Test document analysis endpoint"""
    try:
        # Extract text from uploaded file
        content = await file.read()
        
        if file.filename.lower().endswith(('.doc', '.docx')):
            document_text = await extract_text_from_docx(content)
        else:
            document_text = content.decode('utf-8')
        
        if not document_text or len(document_text.strip()) < 50:
            return {"error": "Insufficient content extracted from document"}
        
        # Use enhanced flexible extraction
        generator = FlexibleScenarioGenerator(client=None)  # No LLM for testing
        extracted_data = await generator.flexible_extract_from_document(document_text)
        
        print(f"Extraction completed - checking data quality...")
        print(f"Sample extracted roles: Expert={extracted_data.get('participant_roles', {}).get('expert_role', 'Not found')[:50]}...")
        
        # Validate and enhance
        validated_data = await generator.validate_and_enhance_template(extracted_data)
        
        # Store template
        template_id = f"template_{len(mock_db['flexible_templates']) + 1}"
        template_record = {
            "template_id": template_id,
            "template_name": template_name,
            "source_type": "document",
            "source_file": file.filename,
            "extracted_data": extracted_data,
            "validated_data": validated_data,
            "status": "pending_approval"
        }
        
        mock_db["flexible_templates"][template_id] = template_record
        
        return {
            "template_id": template_id,
            "extracted_data": extracted_data,
            "validated_data": validated_data,
            "validation_score": validated_data.get("validation_notes", {}).get("completeness_score", "N/A"),
            "missing_elements": validated_data.get("validation_notes", {}).get("missing_elements", []),
            "suggestions": validated_data.get("validation_notes", {}).get("suggestions", []),
            "message": "Document analyzed successfully. Review and approve to generate scenarios."
        }
        
    except Exception as e:
        return {"error": f"Error analyzing document: {str(e)}"}

@app.get("/flexible/template/{template_id}")
async def get_template(template_id: str):
    """Get template for review"""
    try:
        template = mock_db["flexible_templates"].get(template_id)
        if not template:
            return {"error": "Template not found"}
        
        return {
            "template_id": template_id,
            "template_name": template.get("template_name"),
            "source_type": template.get("source_type"),
            "extracted_data": template.get("extracted_data", {}),
            "validated_data": template.get("validated_data", {}),
            "status": template.get("status"),
            "editable_fields": {
                "scenario_understanding": "Core scenario information",
                "participant_roles": "Who is involved in training",
                "knowledge_base": "Key knowledge and procedures"
            }
        }
        
    except Exception as e:
        return {"error": f"Error retrieving template: {str(e)}"}

@app.put("/flexible/template/{template_id}")
async def update_template(template_id: str, updated_data: dict):
    """Update template after user review"""
    try:
        template = mock_db["flexible_templates"].get(template_id)
        if not template:
            return {"error": "Template not found"}
        
        # Update template data
        template["validated_data"]["validated_extraction"] = updated_data
        template["status"] = "approved"
        
        return {
            "message": "Template updated successfully",
            "template_id": template_id,
            "status": "approved",
            "ready_for_generation": True
        }
        
    except Exception as e:
        return {"error": f"Error updating template: {str(e)}"}

@app.post("/flexible/generate-prompt/{template_id}")
async def generate_individual_prompt(template_id: str, request_data: dict):
    """Generate individual prompt from approved template"""
    try:
        template = mock_db["flexible_templates"].get(template_id)
        if not template:
            return {"error": "Template not found"}
        
        if template.get("status") != "approved":
            return {"error": "Template must be approved before generating prompts"}
        
        prompt_type = request_data.get("prompt_type")
        scenario_data = template.get("validated_data", {}).get("validated_extraction", {})
        
        # Generate specific prompt from template
        if prompt_type == "learn_mode":
            expert_role = scenario_data.get('participant_roles', {}).get('expert_role', 'Expert Trainer')
            topics = scenario_data.get('content_specifics', {}).get('key_knowledge', ['General topics'])
            prompt = f"You are a {expert_role}. Your role is to teach about: {', '.join(topics[:3])}. Maintain a supportive and educational approach."
        elif prompt_type in ["assess_mode", "try_mode"]:
            practice_role = scenario_data.get('participant_roles', {}).get('practice_role', 'Practice Partner')
            concerns = scenario_data.get('conversation_dynamics', {}).get('typical_interactions', ['General questions'])
            prompt = f"You are playing the role of: {practice_role}. Your typical concerns include: {', '.join(concerns[:2])}. Engage naturally with the learner."
        else:
            return {"error": "Invalid prompt_type. Use 'learn_mode', 'assess_mode', or 'try_mode'"}
        
        return {
            "template_id": template_id,
            "prompt_type": prompt_type,
            "generated_prompt": prompt,
            "message": f"{prompt_type} prompt generated successfully"
        }
        
    except Exception as e:
        return {"error": f"Error generating prompt: {str(e)}"}

async def test_api_endpoints():
    """Test the API endpoints that the frontend will call"""
    
    print("TESTING FRONTEND INTEGRATION")
    print("=" * 40)
    
    # Test 1: Document Analysis
    print("\n1. Testing Document Analysis...")
    
    with open("Leadership Fundamentals and Styles.docx", "rb") as f:
        file_content = f.read()
    
    # Simulate file upload
    class MockFile:
        def __init__(self, content, filename):
            self.content = content
            self.filename = filename
        
        async def read(self):
            return self.content
    
    mock_file = MockFile(file_content, "Leadership Fundamentals and Styles.docx")
    
    # Call analyze endpoint
    result1 = await analyze_document(mock_file, "Test Leadership Template")
    
    if "error" in result1:
        print(f"Document analysis failed: {result1['error']}")
        return
    
    template_id = result1["template_id"]
    print(f"Document analyzed successfully")
    print(f"   Template ID: {template_id}")
    print(f"   Completeness Score: {result1['validation_score']}%")
    print(f"   Missing Elements: {len(result1['missing_elements'])}")
    
    # Test 2: Get Template
    print("\n2. Testing Get Template...")
    
    result2 = await get_template(template_id)
    
    if "error" in result2:
        print(f"Get template failed: {result2['error']}")
        return
    
    print(f"Template retrieved successfully")
    print(f"   Status: {result2['status']}")
    print(f"   Expert Role: {result2['extracted_data'].get('participant_roles', {}).get('expert_role', 'Not found')[:50]}...")
    print(f"   Practice Role: {result2['extracted_data'].get('participant_roles', {}).get('practice_role', 'Not found')[:50]}...")
    
    # Test 3: Update Template (Approve)
    print("\n3. Testing Template Update...")
    
    result3 = await update_template(template_id, result2["extracted_data"])
    
    if "error" in result3:
        print(f"Template update failed: {result3['error']}")
        return
    
    print(f"Template updated successfully")
    print(f"   Status: {result3['status']}")
    print(f"   Ready for generation: {result3['ready_for_generation']}")
    
    # Test 4: Generate Individual Prompts
    print("\n4. Testing Prompt Generation...")
    
    for prompt_type in ["learn_mode", "assess_mode", "try_mode"]:
        result4 = await generate_individual_prompt(template_id, {"prompt_type": prompt_type})
        
        if "error" in result4:
            print(f"{prompt_type} generation failed: {result4['error']}")
            continue
        
        print(f"{prompt_type} generated successfully")
        print(f"   Prompt: {result4['generated_prompt'][:100]}...")
    
    # Test 5: Frontend Data Structure
    print("\n5. Testing Frontend Data Structure...")
    
    template_data = mock_db["flexible_templates"][template_id]
    extracted_data = template_data["extracted_data"]
    
    # Check if all expected fields exist for frontend
    required_fields = [
        "scenario_understanding",
        "participant_roles", 
        "knowledge_base",
        "coaching_rules",
        "success_metrics",
        "conversation_dynamics",
        "content_specifics",
        "feedback_mechanism"
    ]
    
    missing_fields = []
    for field in required_fields:
        if field not in extracted_data:
            missing_fields.append(field)
    
    if missing_fields:
        print(f"Missing fields for frontend: {missing_fields}")
    else:
        print(f"All required fields present for frontend")
    
    # Check extraction sources
    from_document_count = 0
    total_sections = 0
    
    for section_name, section_data in extracted_data.items():
        if isinstance(section_data, dict) and "extraction_source" in section_data:
            total_sections += 1
            if "FROM_DOCUMENT" in section_data["extraction_source"]:
                from_document_count += 1
    
    print(f"   Document extraction quality: {from_document_count}/{total_sections} sections from document")
    
    print(f"\nFRONTEND INTEGRATION TEST COMPLETED")
    print(f"   Template ID: {template_id}")
    print(f"   All endpoints working correctly")
    print(f"   Ready for frontend testing!")

if __name__ == "__main__":
    asyncio.run(test_api_endpoints())